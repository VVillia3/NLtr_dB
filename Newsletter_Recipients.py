#   Added in version (1.0.5):
#       Duplicate handling now supports households with three or more records that share
#       Address 1 but have different Address 2 values.
#       Manual Save Recipient warnings list every possible matching record instead of only
#       the first database row returned.
#       CSV import checks every possible matching row before deciding whether a complete
#       match should be skipped automatically.
#       When several CSV duplicate candidates exist, the strongest candidate is displayed
#       for review rather than an arbitrary first row.
#       The CSV duplicate-review window now includes Add as Separate New Record so a valid
#       additional household member or address-line variation can be saved without replacing
#       or altering an existing record.
#
#   Added in version (1.0.4):
#       A new Notes field is:
#           Added to the SQLite database schema.
#           Automatically added to an existing database the next time the application
#           launches. No manual database conversion is needed.
#           Displayed beneath the Email List and Mail List checkboxes and above the Save
#           and Reset buttons.
#           Entered in a wrapped, multiline text box with a vertical scrollbar.
#           Saved with new recipients.
#           Loaded for display and editing.
#           Updated when an existing record is changed.
#           Cleared when the form is reset or after a record is saved.
#           Preserved during CSV export and import.
#           Included in Excel exports.
#           Included in CSV duplicate comparisons, selective CSV updates, exception reports,
#           and the Combine Duplicate Records window.
#           Existing CSV backup files that do not contain a notes column will still import
#           normally; imported notes default to blank.
#       Both Avery label exports now use a shared page-building function:
#           The 14-label export creates additional 7-row × 2-column sheets as needed.
#           The 30-label export creates additional 10-row × 3-column sheets as needed.
#           A page break is inserted only when another page is required.
#           The existing label dimensions, margins, ordering, and address formatting
#           remain unchanged.



"""
Newsletter Recipients dB
Copyright (c) 2026 William Tinney

Licensed under the MIT License.
SPDX-License-Identifier: MIT

+AMDG
"""

# MIT License
# 
# Copyright (c) 2026 William Tinney
# 
# Permission is hereby granted, free of charge, to any person obtaining a copy of this 
# software and associated documentation files (the " Software"), to deal in the Software 
# without restriction, including without limitation the rights to use, copy, modify, merge, 
# publish, distribute, sublicense, and/or sell copies of the Software, and to permit persons 
# to whom the Software is furnished to do so, subject to the following conditions:
# 
# The above copyright notice and this permission notice (including the next paragraph) shall 
# be included in all copies or substantial portions of the Software.
# 
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED, 
# INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR 
# PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE 
# FOR ANY CLAIM, DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR 
# OTHERWISE, ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER 
# DEALINGS IN THE SOFTWARE.

import sqlite3
import tkinter as tk
import subprocess
import re
import csv
import os
import getpass
import socket
import sys
from datetime import datetime
from tkinter import messagebox, filedialog, ttk
import tkinter.font as tkfont
from openpyxl import Workbook
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

APP_NAME = "Newsletter Recipients dB"
APP_VERSION = "1.0.5"

def resource_path(filename):
    if hasattr(sys, "_MEIPASS"):
        return os.path.join(sys._MEIPASS, filename)

    return os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        filename
    )

def get_computer_id():
    if os.name == "nt":
        name = getpass.getuser()
    else:
        try:
            name = subprocess.check_output(
                ["scutil", "--get", "ComputerName"],
                text=True
            ).strip()
        except Exception:
            name = socket.gethostname() or getpass.getuser()

    name = re.sub(r"[^A-Za-z0-9_-]", "_", name)
    return name


DB_NAME = f"newsletter_dist-{get_computer_id()}"

if os.name == "nt":
    base_data_dir = os.environ.get(
        "APPDATA",
        os.path.expanduser("~")
    )
    APP_DATA_DIR = os.path.join(
        base_data_dir,
        "NewsletterRecipients"
    )
else:
    APP_DATA_DIR = os.path.join(
        os.path.expanduser("~"),
        "Library",
        "Application Support",
        "NewsletterRecipients"
    )

os.makedirs(APP_DATA_DIR, exist_ok=True)

DB_FILE = os.path.join(APP_DATA_DIR, f"{DB_NAME}.db")

BG_COLOR = "#071D49"
FG_COLOR = "white"

CHECKBOX_BG = BG_COLOR
CHECKBOX_FG = "yellow"

SAVE_BG = "#0B6B4B"
SAVE_FG = "white"
SAVE_ACTIVE_BG = "#14BD85"
SAVE_ACTIVE_FG = "white"

EXPORT_BG = "#0B5B6B"
EXPORT_FG = "white"
EXPORT_ACTIVE_BG = "#13A2BE"

DELETE_RECORD_BG = "#95250F"
DELETE_RECORD_FG = "yellow"
DELETE_RECORD_ACTIVE_BG = "#F9FC00"
DELETE_RECORD_ACTIVE_FG = "black"

SHOW_ALL_RECORDS_BG = "#0B6B4B"
SHOW_ALL_RECORDS_FG = "white"
SHOW_ALL_RECORDS_ACTIVE_BG = "#14BD85"

UPDATE_RECORD_BG = "#0B6B4B"
UPDATE_RECORD_FG = "white"
UPDATE_RECORD_ACTIVE_BG = "#14BD85"

RESET_BG = "#6B0B5B"
RESET_FG = "white"
RESET_ACTIVE_BG = "#BD14A1"

DUPLICATE_DIALOG_BG = "#1D471D"

def get_record_dict_by_id(record_id):
    with sqlite3.connect(DB_FILE) as conn:
        cursor = conn.execute("""
            SELECT id, honorific_title, first_name, last_name,
                   area_code, phone, email,
                   mailaddress_line1, mailaddress_line2,
                   city, state, zip_code,
                   email_list, mail_list, notes
            FROM recipients
            WHERE id = ?
        """, (record_id,))

        row = cursor.fetchone()

    if not row:
        return None

    fields = [
        "id",
        "honorific_title",
        "first_name",
        "last_name",
        "area_code",
        "phone",
        "email",
        "mailaddress_line1",
        "mailaddress_line2",
        "city",
        "state",
        "zip_code",
        "email_list",
        "mail_list",
        "notes"
    ]

    return dict(zip(fields, row))
    
def combine_selected_duplicates():
    selected_items = search_results_box.selection()

    if len(selected_items) != 2:
        show_error(
            "Select Two Records",
            "Select exactly two records in the Search Results table before combining duplicates."
        )
        return

    selected_ids = []

    for item in selected_items:
        values = search_results_box.item(item, "values")

        if not values or not str(values[0]).isdigit():
            show_error(
                "Invalid Selection",
                "One selected search result does not contain a valid record ID."
            )
            return

        selected_ids.append(values[0])

    record1 = get_record_dict_by_id(selected_ids[0])
    record2 = get_record_dict_by_id(selected_ids[1])

    if not record1 or not record2:
        show_error(
            "Record Not Found",
            "One or both selected records could not be found in the database."
        )
        return

    dialog = tk.Toplevel(root)
    dialog.withdraw()
    dialog.title("Combine Duplicate Records")
    dialog.configure(bg=DUPLICATE_DIALOG_BG)
    dialog.transient(root)

    result = {
        "saved": False
    }

    field_labels = [
        ("honorific_title", "Title"),
        ("first_name", "First Name"),
        ("last_name", "Last Name"),
        ("area_code", "Area Code"),
        ("phone", "Phone"),
        ("email", "Email"),
        ("mailaddress_line1", "Address 1"),
        ("mailaddress_line2", "Address 2"),
        ("city", "City"),
        ("state", "State"),
        ("zip_code", "ZIP Code"),
        ("email_list", "Email List"),
        ("mail_list", "Mail List"),
        ("notes", "Notes")
    ]

    warning_text = (
        "WARNING: This operation will delete both original records and save one new combined record "
        "with a new unique ID. Review all fields carefully before saving."
    )

    instruction_text = (
        "Instructions: Use the R1 and R2 buttons to copy values from Record 1 or Record 2 into the "
        "Combined Value column, or type your own corrected value directly into the Combined Value field. "
        "The Combined Value fields will be saved exactly as shown; normal capitalization and punctuation "
        "rules are not applied in this window. For Email List and Mail List, use 1 for selected/on and "
        "0 for unselected/off."
    )

    tk.Label(
        dialog,
        text=warning_text,
        bg=DUPLICATE_DIALOG_BG,
        fg="#F9FC00",
        font=("TkDefaultFont", 12, "bold"),
        wraplength=900,
        justify="left"
    ).grid(row=0, column=0, columnspan=5, sticky="w", padx=10, pady=(10, 6))

    tk.Label(
        dialog,
        text=instruction_text,
        bg=DUPLICATE_DIALOG_BG,
        fg=FG_COLOR,
        wraplength=900,
        justify="left"
    ).grid(row=1, column=0, columnspan=5, sticky="w", padx=10, pady=(0, 12))

    tk.Label(dialog, text="FIELD", bg=DUPLICATE_DIALOG_BG, fg=FG_COLOR).grid(
        row=2, column=0, sticky="w", padx=10, pady=2
    )

    tk.Label(
        dialog,
        text=f"RECORD 1 (ID {record1['id']})",
        bg=DUPLICATE_DIALOG_BG,
        fg=FG_COLOR
    ).grid(
        row=2, column=1, sticky="w", padx=10, pady=2
    )

    tk.Label(
        dialog,
        text=f"RECORD 2 (ID {record2['id']})",
        bg=DUPLICATE_DIALOG_BG,
        fg=FG_COLOR
    ).grid(
        row=2,
        column=2,
        sticky="w",
        padx=10, 
        pady=2
    )
    
    tk.Label(
        dialog,
        text="USE",
        bg=DUPLICATE_DIALOG_BG,
        fg=FG_COLOR,
        anchor="center"
    ).grid(
        row=2,
        column=3,
        sticky="ew",
        padx=10,
        pady=2
    )
    
    tk.Label(
        dialog,
        text="COMBINED VALUE",
        bg=DUPLICATE_DIALOG_BG,
        fg=FG_COLOR
    ).grid(
        row=2,
        column=4,
        sticky="ew",
        padx=10,
        pady=2
    )
        
    combined_vars = {}
    combined_widgets = {}
    source_buttons = {}

    SOURCE_BUTTON_BG = "#153315"

    def get_combined_value(field_name):
        if field_name == "notes":
            return combined_widgets[field_name].get("1.0", "end-1c")

        return combined_vars[field_name].get()

    def set_combined_value(field_name, value):
        if field_name == "notes":
            widget = combined_widgets[field_name]
            widget.delete("1.0", tk.END)
            widget.insert("1.0", value)
            widget.edit_modified(False)
            update_source_button_highlights()
            return

        combined_vars[field_name].set(value)

    def on_notes_modified(event):
        if event.widget.edit_modified():
            update_source_button_highlights()
            event.widget.edit_modified(False)

    for index, (field_name, label_text) in enumerate(field_labels, start=3):
        value1 = "" if record1.get(field_name) is None else str(record1.get(field_name))
        value2 = "" if record2.get(field_name) is None else str(record2.get(field_name))

        display_value1 = value1
        display_value2 = value2

        if field_name == "area_code":
            display_value1 = format_area_code(value1)
            display_value2 = format_area_code(value2)

        if field_name == "phone":
            display_value1 = format_phone(value1)
            display_value2 = format_phone(value2)

        default_value = value1 if value1 else value2

        field_name_color = "#FC0000" if value1.strip() != value2.strip() else FG_COLOR

        tk.Label(
            dialog,
            text=label_text,
            bg=DUPLICATE_DIALOG_BG,
            fg=field_name_color
        ).grid(row=index, column=0, sticky="w", padx=10, pady=2)

        tk.Label(
            dialog,
            text=display_value1,
            bg=DUPLICATE_DIALOG_BG,
            fg="aqua",
            anchor="w",
            justify="left",
            wraplength=220
        ).grid(row=index, column=1, sticky="w", padx=10, pady=2)

        tk.Label(
            dialog,
            text=display_value2,
            bg=DUPLICATE_DIALOG_BG,
            fg="lime",
            anchor="w",
            justify="left",
            wraplength=220
        ).grid(row=index, column=2, sticky="w", padx=10, pady=2)

        if field_name == "notes":
            notes_frame = tk.Frame(dialog)
            notes_frame.grid(row=index, column=4, sticky="w", padx=10, pady=2)

            notes_scrollbar = tk.Scrollbar(
                notes_frame,
                orient="vertical"
            )
            notes_scrollbar.pack(side="right", fill="y")

            notes_widget = tk.Text(
                notes_frame,
                width=32,
                height=4,
                wrap="word",
                yscrollcommand=notes_scrollbar.set
            )
            notes_widget.pack(side="left", fill="both", expand=True)
            notes_scrollbar.config(command=notes_widget.yview)

            notes_widget.insert("1.0", default_value)
            notes_widget.edit_modified(False)
            notes_widget.bind("<<Modified>>", on_notes_modified)

            combined_widgets[field_name] = notes_widget
        else:
            var = tk.StringVar(value=default_value)
            combined_vars[field_name] = var

            entry = tk.Entry(
                dialog,
                textvariable=var,
                width=32
            )
            entry.grid(row=index, column=4, sticky="w", padx=10, pady=2)
                
        button_frame = tk.Frame(
            dialog,
            bg=DUPLICATE_DIALOG_BG
        )
        button_frame.grid(
            row=index,
            column=3,
            sticky="",
            padx=10,
            pady=2
        )

        r1_button = ColorButton(
            button_frame,
            text="R1",
            width=4,
            bg=SOURCE_BUTTON_BG,
            fg="aqua",
            activebackground=EXPORT_ACTIVE_BG,
            activeforeground=EXPORT_FG,
            command=lambda f=field_name, v=value1: set_combined_value(f, v)
        )
        r1_button.grid(row=0, column=0, padx=(0, 4))

        r2_button = ColorButton(
            button_frame,
            text="R2",
            width=4,
            bg=SOURCE_BUTTON_BG,
            fg="lime",
            activebackground=SHOW_ALL_RECORDS_ACTIVE_BG,
            activeforeground=SHOW_ALL_RECORDS_FG,
            command=lambda f=field_name, v=value2: set_combined_value(f, v)
        )
        r2_button.grid(row=0, column=1)

        source_buttons[field_name] = {
            "r1_button": r1_button,
            "r2_button": r2_button,
            "value1": value1,
            "value2": value2
        }

    def update_source_button_highlights(*args):
        for field_name, button_data in source_buttons.items():
            current_value = get_combined_value(field_name)

            r1_button = button_data["r1_button"]
            r2_button = button_data["r2_button"]
            value1 = button_data["value1"]
            value2 = button_data["value2"]

            if current_value == value1:
                r1_button.config(
                    bg=EXPORT_ACTIVE_BG,
                    fg=EXPORT_FG
                )
            else:
                r1_button.config(
                    bg=SOURCE_BUTTON_BG,
                    fg="aqua"
                )

            if current_value == value2:
                r2_button.config(
                    bg=SHOW_ALL_RECORDS_ACTIVE_BG,
                    fg=SHOW_ALL_RECORDS_FG
                )
            else:
                r2_button.config(
                    bg=SOURCE_BUTTON_BG,
                    fg="lime"
                )

    for var in combined_vars.values():
        var.trace_add("write", update_source_button_highlights)

    update_source_button_highlights()

    button_row = len(field_labels) + 3

    def save_combined_record():
        combined_record = {
            field_name: get_combined_value(field_name).strip()
            for field_name, label_text in field_labels
        }

        combined_record["email_list"] = csv_int(combined_record["email_list"])
        combined_record["mail_list"] = csv_int(combined_record["mail_list"])

        if not combined_record["first_name"] or not combined_record["last_name"]:
            show_error(
                "Missing Information",
                "The combined record must include First Name and Last Name."
            )
            return

        if not combined_record["email_list"] and not combined_record["mail_list"]:
            show_error(
                "List Selection Required",
                "The combined record must be assigned to Email List, Mail List, or both."
            )
            return

        if combined_record["email_list"] and not combined_record["email"]:
            show_error(
                "Missing Email",
                "An email address is required if Email List is selected."
            )
            return

        if combined_record["mail_list"] and (
            not combined_record["mailaddress_line1"]
            or not combined_record["city"]
            or not combined_record["state"]
            or not combined_record["zip_code"]
        ):
            show_error(
                "Missing Mailing Address",
                "Address 1, City, State, and ZIP Code are required if Mail List is selected."
            )
            return

        confirm = ask_yes_no(
            "Confirm Combine Duplicates",
            f"This will delete original record IDs {record1['id']} and {record2['id']} "
            f"and save one new combined record with a new unique ID.\n\n"
            f"This cannot be undone.\n\n"
            f"Continue?"
        )

        if not confirm:
            return

        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.execute("""
                INSERT INTO recipients (
                    honorific_title, first_name, last_name,
                    area_code, phone, email,
                    mailaddress_line1, mailaddress_line2,
                    city, state, zip_code,
                    email_list, mail_list, notes
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                combined_record["honorific_title"],
                combined_record["first_name"],
                combined_record["last_name"],
                combined_record["area_code"],
                combined_record["phone"],
                combined_record["email"],
                combined_record["mailaddress_line1"],
                combined_record["mailaddress_line2"],
                combined_record["city"],
                combined_record["state"],
                combined_record["zip_code"],
                combined_record["email_list"],
                combined_record["mail_list"],
                combined_record["notes"]
            ))

            new_id = cursor.lastrowid

            conn.execute(
                "DELETE FROM recipients WHERE id IN (?, ?)",
                (record1["id"], record2["id"])
            )

        result["saved"] = True
        dialog.destroy()

        mark_unbacked_changes()
        refresh_database_status(update_timestamp=True)
        show_all_records_inline()
        clear_form()

        show_info(
            "Duplicates Combined",
            f"Original record IDs {record1['id']} and {record2['id']} were deleted.\n\n"
            f"New combined record ID: {new_id}"
        )

    def cancel_combine():
        dialog.destroy()

    ColorButton(
        dialog,
        text="Save Combined / Delete Originals",
        bg=DELETE_RECORD_BG,
        fg="yellow",
        activebackground=DELETE_RECORD_ACTIVE_BG,
        activeforeground="black",
        width=30,
        command=save_combined_record
    ).grid(
        row=button_row,
        column=4,
        sticky="",
        padx=10,
        pady=(15,5)
    )

    ColorButton(
        dialog,
        text="Cancel",
        bg=RESET_BG,
        fg=RESET_FG,
        activebackground=RESET_ACTIVE_BG,
        activeforeground=RESET_FG,
        width=18,
        command=cancel_combine
    ).grid(
        row=button_row + 1,
        column=4,
        sticky="",
        padx=10,
        pady=(5, 10)
    )
    
    center_toplevel_over_parent(dialog, root)
    dialog.grab_set()
    dialog.wait_window()
    

class ColorButton(tk.Label):
    def __init__(
        self,
        master,
        text="",
        command=None,
        width=24,
        bg="#00D5FF",
        fg="black",
        activebackground=None,
        activeforeground=None,
        state="normal",
        **kwargs
    ):
        self.command = command
        self.normal_bg = bg
        self.normal_fg = fg
        self.active_bg = activebackground or bg
        self.active_fg = activeforeground or fg
        self.disabled_bg = "#BFBFBF"
        self.disabled_fg = "#777777"
        self._state = state

        super().__init__(
            master,
            text=text,
            width=width,
            bg=bg,
            fg=fg,
            padx=8,
            pady=4,
            relief="flat",
            borderwidth=0,
            anchor="center",
            cursor="hand2",
            **kwargs
        )

        self.bind("<Button-1>", self._on_click)
        self.bind("<Enter>", self._on_enter)
        self.bind("<Leave>", self._on_leave)

        self.config(state=state)

    def _on_click(self, event=None):
        if self._state != "disabled" and self.command:
            self.command()

    def _on_enter(self, event=None):
        if self._state != "disabled":
            super().config(bg=self.active_bg, fg=self.active_fg)

    def _on_leave(self, event=None):
        if self._state != "disabled":
            super().config(bg=self.normal_bg, fg=self.normal_fg)

    def config(self, cnf=None, **kwargs):
        if cnf:
            kwargs.update(cnf)

        if "command" in kwargs:
            self.command = kwargs.pop("command")

        if "state" in kwargs:
            self._state = kwargs.pop("state")

        if "bg" in kwargs:
            self.normal_bg = kwargs["bg"]

        if "fg" in kwargs:
            self.normal_fg = kwargs["fg"]

        if "activebackground" in kwargs:
            self.active_bg = kwargs.pop("activebackground")

        if "activeforeground" in kwargs:
            self.active_fg = kwargs.pop("activeforeground")

        kwargs.pop("highlightbackground", None)
        kwargs.pop("highlightcolor", None)
        kwargs.pop("highlightthickness", None)
        kwargs.pop("borderwidth", None)
        kwargs.pop("relief", None)
        kwargs.pop("overrelief", None)

        if self._state == "disabled":
            kwargs["bg"] = self.disabled_bg
            kwargs["fg"] = self.disabled_fg
            kwargs["cursor"] = ""
        else:
            kwargs["bg"] = self.normal_bg
            kwargs["fg"] = self.normal_fg
            kwargs["cursor"] = "hand2"

        super().config(**kwargs)

    configure = config

def style_button(
    button,
    bg=EXPORT_BG,
    fg="black",
    active_bg=None,
    active_fg=None,
    width=24
):
    if active_bg is None:
        active_bg = bg

    if active_fg is None:
        active_fg = fg

    button.config(
        width=width,
        bg=bg,
        fg=fg,
        activebackground=active_bg,
        activeforeground=active_fg,
        highlightbackground=bg,
        highlightcolor=bg,
        highlightthickness=0,
        borderwidth=0,
        relief="flat",
        overrelief="raised"
    )

CURRENT_RECORD_ID = None

SESSION_HAS_UNBACKED_CHANGES = False


def mark_unbacked_changes():
    global SESSION_HAS_UNBACKED_CHANGES
    SESSION_HAS_UNBACKED_CHANGES = True


def mark_csv_backup_complete():
    global SESSION_HAS_UNBACKED_CHANGES
    SESSION_HAS_UNBACKED_CHANGES = False
    

def normalize_address(address, lowercase_and=False, preserve_exact=False):
    if preserve_exact:
        return address.strip()

    address = address.strip().title()

    if lowercase_and:
        address = re.sub(r"\bAnd\b", "and", address)

    address = address.replace("Po Box", "PO Box")

    address = re.sub(r'(\d+)St\b', r'\1st', address)
    address = re.sub(r'(\d+)Nd\b', r'\1nd', address)
    address = re.sub(r'(\d+)Rd\b', r'\1rd', address)
    address = re.sub(r'(\d+)Th\b', r'\1th', address)

    address = re.sub(r'\bNe\b', 'NE', address)
    address = re.sub(r'\bNw\b', 'NW', address)
    address = re.sub(r'\bSe\b', 'SE', address)
    address = re.sub(r'\bSw\b', 'SW', address)

    return address

def normalize_city(city, lowercase_and=True, preserve_exact=False):
    if preserve_exact:
        return city.strip()

    city = city.strip().title()

    if lowercase_and:
        city = re.sub(r"\bAnd\b", "and", city)

    return city

def open_file(file_path):
    if os.name == "nt":
        os.startfile(file_path)
    else:
        subprocess.run(["open", file_path])
        
def open_user_manual():
    manual_path = resource_path(
        "Newsletter_Recipients_User_Manual.pdf"
    )

    if not os.path.exists(manual_path):
        show_error(
            "User Manual Not Found",
            "The bundled user manual could not be located."
        )
        return

    open_file(manual_path)

def normalize_name(
    name,
    is_title=False,
    lowercase_and=False,
    preserve_exact=False
):
    if preserve_exact:
        return name.strip()

    name = name.strip().title()

    if lowercase_and:
        name = re.sub(r"\bAnd\b", "and", name)

    if is_title and name:
        no_period_titles = {
            "Miss"
        }

        if name not in no_period_titles and not name.endswith("."):
            name += "."

    return name
    
def digits_only(value):
    return re.sub(r"\D", "", value.strip())


def format_area_code(area_code):
    return f"({area_code})" if area_code else ""


def format_phone(phone):
    return f"{phone[:3]}-{phone[3:]}" if phone else ""


def get_form_data_or_error():
    hon_title = normalize_name(
        honorific_title_entry.get(),
        is_title=True,
        preserve_exact=preserve_title_var.get()
    )

    first = normalize_name(
        first_name_entry.get(),
        lowercase_and=True,
        preserve_exact=preserve_first_name_var.get()
    )

    last = normalize_name(
        last_name_entry.get(),
        preserve_exact=preserve_last_name_var.get()
    )

    areacode = digits_only(area_code_entry.get())
    phone = digits_only(phone_entry.get())
    email = email_entry.get().strip().lower()

    address1 = normalize_address(
        mailaddress_line1_entry.get(),
        lowercase_and=True,
        preserve_exact=preserve_address1_var.get()
    )

    address2 = normalize_address(
        mailaddress_line2_entry.get(),
        lowercase_and=True,
        preserve_exact=preserve_address2_var.get()
    )

    city = normalize_city(
        city_entry.get(),
        lowercase_and=True,
        preserve_exact=preserve_city_var.get()
    )

    state = state_entry.get().strip().upper()
    zip_code = zip_code_entry.get().strip()
    listemail = email_list_var.get()
    listmail = mail_list_var.get()
    notes = notes_text.get("1.0", tk.END).strip()

    if not first or not last:
        show_error(
            "Error",
            "Missing Information.\n\nFirst and last name are required."
        )
        return None

    if not listemail and not listmail:
        show_error(
            "Error",
            "List selection required.\n\nSelect Email List or Mail List."
        )
        return None

    if areacode and (not areacode.isdigit() or len(areacode) != 3):
        show_error(
            "Error",
            "Invalid area code.\n\nEnter 3 digits, with or without parentheses."
        )
        return None

    if phone and (not phone.isdigit() or len(phone) != 7):
        show_error(
            "Error",
            "Invalid phone number.\n\nEnter 7 digits only, with or without a dash."
        )
        return None

    if listemail and not email:
        show_error(
            "Error",
            "Missing email address.\n\nAn email address is required if Email List is selected."
        )
        return None

    if email and ("@" not in email or "." not in email.split("@")[-1]):
        show_error(
            "Error",
            "Invalid Email.\n\nPlease enter a complete email address."
        )
        return None

    if listmail and (not address1 or not city or not state or not zip_code):
        show_error(
            "Error",
            "Missing complete mailing address.\n\nIf Mail List is selected, confirm all entries are completed."
        )
        return None

    if zip_code and (not zip_code.isdigit() or len(zip_code) != 5):
        show_error(
            "Error",
            "Invalid ZIP code.\n\nZIP code must contain 5 digits."
        )
        return None

    return (
        hon_title, first, last, areacode, phone, email,
        address1, address2, city, state, zip_code, listemail, listmail, notes
    )

def setup_database():
    with sqlite3.connect(DB_FILE) as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS recipients (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                honorific_title TEXT,
                first_name TEXT NOT NULL,
                last_name TEXT NOT NULL,
                area_code TEXT,
                phone TEXT,
                email TEXT,
                mailaddress_line1 TEXT,
                mailaddress_line2 TEXT,
                city TEXT,
                state TEXT,
                zip_code TEXT,
                email_list INTEGER NOT NULL DEFAULT 0,
                mail_list INTEGER NOT NULL DEFAULT 0,
                notes TEXT
            )
        """)

        existing_columns = [
            row[1]
            for row in conn.execute("PRAGMA table_info(recipients)")
        ]

        if "mailaddress_line2" not in existing_columns:
            conn.execute("""
                ALTER TABLE recipients
                ADD COLUMN mailaddress_line2 TEXT
            """)
            
        if "notes" not in existing_columns:
            conn.execute("""
                ALTER TABLE recipients
                ADD COLUMN notes TEXT
            """)
            
        conn.execute("""
            CREATE INDEX IF NOT EXISTS idx_last_name
            ON recipients(last_name)
        """)

        conn.execute("""
            CREATE INDEX IF NOT EXISTS idx_email
            ON recipients(email)
        """)

        conn.execute("""
            CREATE INDEX IF NOT EXISTS idx_phone
            ON recipients(phone)
        """)

        conn.execute("""
            CREATE INDEX IF NOT EXISTS idx_address
            ON recipients(mailaddress_line1)
        """)
        
        conn.execute("""
            CREATE TABLE IF NOT EXISTS app_metadata (
                key TEXT PRIMARY KEY,
                value TEXT
            )
        """)

def get_saved_last_update():
    with sqlite3.connect(DB_FILE) as conn:
        row = conn.execute("""
            SELECT value
            FROM app_metadata
            WHERE key = 'last_update'
        """).fetchone()

    return row[0] if row else ""


def save_last_update():
    timestamp = datetime.now().strftime("%H:%M:%S %a %b %d %Y")

    with sqlite3.connect(DB_FILE) as conn:
        conn.execute("""
            INSERT INTO app_metadata (key, value)
            VALUES ('last_update', ?)
            ON CONFLICT(key) DO UPDATE SET value = excluded.value
        """, (timestamp,))

    return timestamp

def save_recipient():
    form_data = get_form_data_or_error()

    if form_data is None:
        return

    (
        hon_title, first, last, areacode, phone, email,
        address1, address2, city, state, zip_code, listemail, listmail, notes
    ) = form_data

    with sqlite3.connect(DB_FILE) as conn:
        cursor = conn.execute("""
            SELECT id, first_name, last_name,
                   mailaddress_line1, mailaddress_line2,
                   email, phone
            FROM recipients
            WHERE (
                    lower(email) = lower(?)
                    AND email != ''
                )
            OR (
                phone = ?
                AND phone != ''
                )
            OR (
                lower(last_name) = lower(?)
                AND lower(mailaddress_line1) = lower(?)
                AND last_name != ''
                AND mailaddress_line1 != ''
                )
            ORDER BY id
        """, (email, phone, last, address1))

        duplicates = cursor.fetchall()

        if duplicates:
            existing_record_text = []

            for (
                existing_id, ex_first, ex_last,
                ex_address1, ex_address2,
                ex_email, ex_phone
            ) in duplicates:
                address_lines = "\n".join(
                    line
                    for line in [ex_address1, ex_address2]
                    if line
                )

                existing_record_text.append(
                    f"Existing ID: {existing_id}\n"
                    f"{ex_first} {ex_last}\n"
                    f"{address_lines}\n"
                    f"{ex_email}\n"
                    f"{format_phone(ex_phone)}"
                )

            proceed = ask_yes_no(
                "Possible Duplicate",
                f"Possible duplicates found:\n\n"
                f"{'\n\n'.join(existing_record_text)}\n\n"
                f"New record:\n"
                f"{first} {last}\n"
                f"{address1}\n"
                f"{address2}\n"
                f"{email}\n"
                f"{format_phone(phone)}\n\n"
                f"Do you still want to save this as a new record?"
            )

            if not proceed:
                return

        conn.execute("""
            INSERT INTO recipients (
                honorific_title, first_name, last_name, area_code, phone, email,
                mailaddress_line1, mailaddress_line2, city, state, zip_code, email_list,
                mail_list, notes
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, form_data)

    mark_unbacked_changes()
    refresh_database_status(update_timestamp=True)

    honorific_title_entry.delete(0, tk.END)
    first_name_entry.delete(0, tk.END)
    last_name_entry.delete(0, tk.END)
    area_code_entry.delete(0, tk.END)
    phone_entry.delete(0, tk.END)
    email_entry.delete(0, tk.END)
    mailaddress_line1_entry.delete(0, tk.END)
    mailaddress_line2_entry.delete(0, tk.END)
    city_entry.delete(0, tk.END)
    state_entry.delete(0, tk.END)
    zip_code_entry.delete(0, tk.END)
    notes_text.delete("1.0", tk.END)
    email_list_var.set(0)
    mail_list_var.set(0)

    preserve_title_var.set(0)
    preserve_first_name_var.set(0)
    preserve_last_name_var.set(0)
    preserve_address1_var.set(0)
    preserve_address2_var.set(0)
    preserve_city_var.set(0)

    refresh_current_search_results()

    honorific_title_entry.focus_set()

    show_info(
        "Saved",
        "Recipient saved."
    )

def clear_form():
    honorific_title_entry.delete(0, tk.END)
    first_name_entry.delete(0, tk.END)
    last_name_entry.delete(0, tk.END)
    area_code_entry.delete(0, tk.END)
    phone_entry.delete(0, tk.END)
    email_entry.delete(0, tk.END)
    mailaddress_line1_entry.delete(0, tk.END)
    mailaddress_line2_entry.delete(0, tk.END)
    city_entry.delete(0, tk.END)
    state_entry.delete(0, tk.END)
    zip_code_entry.delete(0, tk.END)
    notes_text.delete("1.0", tk.END)

    email_list_var.set(0)
    mail_list_var.set(0)
    
    preserve_title_var.set(0)
    preserve_first_name_var.set(0)
    preserve_last_name_var.set(0)
    preserve_address1_var.set(0)
    preserve_address2_var.set(0)
    preserve_city_var.set(0)

    edit_id_entry.delete(0, tk.END)

    global CURRENT_RECORD_ID
    CURRENT_RECORD_ID = None

    save_button.config(text="Save Recipient")
    save_button.config(state="normal")
    delete_button.config(state="disabled")

def export_xlsx(query, default_filename):
    file_path = filedialog.asksaveasfilename(
        parent=root,
        defaultextension=".xlsx",
        initialfile=default_filename,
        filetypes=[("Excel files", "*.xlsx")]
    )

    if not file_path:
        return

    with sqlite3.connect(DB_FILE) as conn:
        cursor = conn.execute(query)
        rows = cursor.fetchall()
        headers = [description[0] for description in cursor.description]

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Recipients"

    sheet.append(headers)

    area_code_index = headers.index("area_code") if "area_code" in headers else None
    phone_index = headers.index("phone") if "phone" in headers else None

    for row in rows:
        row = list(row)

        if area_code_index is not None:
            row[area_code_index] = format_area_code(row[area_code_index])

        if phone_index is not None:
            row[phone_index] = format_phone(row[phone_index])

        sheet.append(row)

    workbook.save(file_path)

    open_file(file_path)

    show_info(
        "Export Complete",
        f"Excel file saved:\n\n{file_path}"
    )

def export_csv_all():
    file_path = filedialog.asksaveasfilename(
        parent=root,
        defaultextension=".csv",
        initialfile=f"{DB_NAME}_{datetime.now().strftime('%Y-%m-%d')}.csv",
        filetypes=[("CSV files", "*.csv")]
    )

    if not file_path:
        return

    with sqlite3.connect(DB_FILE) as conn:
        cursor = conn.execute("""
            SELECT
                ? AS source_db,
                honorific_title,
                first_name,
                last_name,
                area_code,
                phone,
                email,
                mailaddress_line1,
                mailaddress_line2,
                city,
                state,
                zip_code,
                email_list,
                mail_list,
                notes
            FROM recipients
            ORDER BY lower(last_name), lower(first_name)
        """, (DB_NAME,))

        rows = cursor.fetchall()
        headers = [description[0] for description in cursor.description]

    with open(file_path, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(headers)
        writer.writerows(rows)

    mark_csv_backup_complete()

    open_file(file_path)

    show_info(
        "Export Complete",
        f"CSV file saved:\n\n{file_path}"
    )

CSV_RECIPIENT_FIELDS = [
    "honorific_title",
    "first_name",
    "last_name",
    "area_code",
    "phone",
    "email",
    "mailaddress_line1",
    "mailaddress_line2",
    "city",
    "state",
    "zip_code",
    "email_list",
    "mail_list",
    "notes"
]


def csv_record_matches_existing_record(existing_record, csv_record):
    for field_name in CSV_RECIPIENT_FIELDS:
        existing_value = existing_record.get(field_name, "")
        csv_value = csv_record.get(field_name, "")

        if field_name in ("email_list", "mail_list"):
            if int(existing_value or 0) != int(csv_value or 0):
                return False
        else:
            if str(existing_value or "").strip() != str(csv_value or "").strip():
                return False

    return True

def csv_int(value):
    value = str(value).strip()

    if value in ("1", "true", "True", "yes", "Yes", "Y", "y"):
        return 1

    return 0


def csv_compare_text(value):
    return str(value or "").strip().casefold()


def csv_duplicate_match_score(existing_record, csv_record):
    """
    Rank possible duplicate candidates so the review window opens against
    the most likely matching database record when more than one row matches.
    """
    score = 0

    existing_email = csv_compare_text(existing_record.get("email"))
    csv_email = csv_compare_text(csv_record.get("email"))

    if csv_email and existing_email == csv_email:
        score += 100

    existing_phone = csv_compare_text(existing_record.get("phone"))
    csv_phone = csv_compare_text(csv_record.get("phone"))

    if csv_phone and existing_phone == csv_phone:
        score += 100

    if (
        csv_compare_text(existing_record.get("last_name"))
        == csv_compare_text(csv_record.get("last_name"))
    ):
        score += 20

    existing_addresses = {
        csv_compare_text(existing_record.get("mailaddress_line1")),
        csv_compare_text(existing_record.get("mailaddress_line2"))
    }

    csv_addresses = {
        csv_compare_text(csv_record.get("mailaddress_line1")),
        csv_compare_text(csv_record.get("mailaddress_line2"))
    }

    existing_addresses.discard("")
    csv_addresses.discard("")

    score += 10 * len(existing_addresses & csv_addresses)

    if (
        csv_compare_text(existing_record.get("mailaddress_line1"))
        == csv_compare_text(csv_record.get("mailaddress_line1"))
        and csv_compare_text(existing_record.get("mailaddress_line2"))
        == csv_compare_text(csv_record.get("mailaddress_line2"))
    ):
        score += 40

    return score


def database_duplicate_row_to_record(duplicate):
    (
        existing_id, ex_hon_title, ex_first, ex_last,
        ex_area_code, ex_phone, ex_email, ex_address1,
        ex_address2, ex_city, ex_state, ex_zip_code,
        ex_email_list, ex_mail_list, ex_notes
    ) = duplicate

    existing_record = {
        "honorific_title": ex_hon_title or "",
        "first_name": ex_first or "",
        "last_name": ex_last or "",
        "area_code": ex_area_code or "",
        "phone": ex_phone or "",
        "email": ex_email or "",
        "mailaddress_line1": ex_address1 or "",
        "mailaddress_line2": ex_address2 or "",
        "city": ex_city or "",
        "state": ex_state or "",
        "zip_code": ex_zip_code or "",
        "email_list": ex_email_list or 0,
        "mail_list": ex_mail_list or 0,
        "notes": ex_notes or ""
    }

    return existing_id, existing_record


def insert_csv_record(conn, csv_record):
    conn.execute("""
        INSERT INTO recipients (
            honorific_title, first_name, last_name,
            area_code, phone, email,
            mailaddress_line1, mailaddress_line2, city, state, zip_code,
            email_list, mail_list, notes
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        csv_record["honorific_title"],
        csv_record["first_name"],
        csv_record["last_name"],
        csv_record["area_code"],
        csv_record["phone"],
        csv_record["email"],
        csv_record["mailaddress_line1"],
        csv_record["mailaddress_line2"],
        csv_record["city"],
        csv_record["state"],
        csv_record["zip_code"],
        csv_record["email_list"],
        csv_record["mail_list"],
        csv_record["notes"]
    ))


def choose_csv_fields_for_update(existing_record, csv_record):
    dialog = tk.Toplevel(root)
    dialog.withdraw()
    dialog.title("Review Possible CSV Duplicate")
    dialog.configure(bg=DUPLICATE_DIALOG_BG)
    dialog.resizable(False, False)
    dialog.transient(root)

    result = {
        "action": None,
        "selected_fields": []
        }

    field_labels = [
        ("honorific_title", "Title"),
        ("first_name", "First Name"),
        ("last_name", "Last Name"),
        ("area_code", "Area Code"),
        ("phone", "Phone"),
        ("email", "Email"),
        ("mailaddress_line1", "Address 1"),
        ("mailaddress_line2", "Address 2"),
        ("city", "City"),
        ("state", "State"),
        ("zip_code", "ZIP Code"),
        ("email_list", "Email List"),
        ("mail_list", "Mail List"),
        ("notes", "Notes")
    ]

    tk.Label(
        dialog,
        text="Possible duplicate found. Compare the existing database record with the CSV record.",
        bg=DUPLICATE_DIALOG_BG,
        fg=FG_COLOR,
        font=("TkDefaultFont", 12, "bold")
    ).grid(row=0, column=0, columnspan=5, sticky="w", padx=10, pady=(10, 8))

    tk.Label(dialog, text="Field", bg=DUPLICATE_DIALOG_BG, fg=FG_COLOR).grid(
        row=1, column=0, sticky="w", padx=10, pady=2
    )

    tk.Label(
        dialog,
        text="Existing Database Record",
        bg=DUPLICATE_DIALOG_BG,
        fg=FG_COLOR
    ).grid(
        row=1, column=1, sticky="w", padx=10, pady=2
    )

    tk.Label(dialog, text="Use CSV?", bg=DUPLICATE_DIALOG_BG, fg=FG_COLOR).grid(
        row=1, column=2, sticky="", padx=10, pady=2
    )
    tk.Label(dialog, text="CSV Record", bg=DUPLICATE_DIALOG_BG, fg=FG_COLOR).grid(
        row=1, column=3, sticky="w", padx=10, pady=2
    )

    selected_vars = {}
    csv_buttons = {}

    CSV_BUTTON_NORMAL_BG = "#153315"
    CSV_BUTTON_NORMAL_FG = "aqua"

    def update_action_button_states():
        any_csv_selected = any(
            var.get() == 1
            for var in selected_vars.values()
        )

        if any_csv_selected:
            replace_all_button.config(state="disabled")
            update_selected_button.config(state="normal")
        else:
            replace_all_button.config(state="normal")
            update_selected_button.config(state="disabled")
            
    def toggle_csv_field(field_name):
        var = selected_vars[field_name]
        button = csv_buttons[field_name]

        if var.get() == 1:
            var.set(0)
            button.config(
                bg=CSV_BUTTON_NORMAL_BG,
                fg=CSV_BUTTON_NORMAL_FG
            )
        else:
            var.set(1)
            button.config(
                bg=EXPORT_ACTIVE_BG,
                fg=EXPORT_FG
            )

        update_action_button_states()
        
    for index, (field_name, label_text) in enumerate(field_labels, start=2):
        existing_value = existing_record.get(field_name, "")
        csv_value = csv_record.get(field_name, "")

        existing_compare = str(existing_value).strip()
        csv_compare = str(csv_value).strip()

        field_name_color = "#FC0000" if existing_compare != csv_compare else FG_COLOR

        if field_name == "area_code":
            existing_value = format_area_code(existing_value)
            csv_value = format_area_code(csv_value)

        if field_name == "phone":
            existing_value = format_phone(existing_value)
            csv_value = format_phone(csv_value)

        var = tk.IntVar(value=0)
        selected_vars[field_name] = var

        tk.Label(
            dialog,
            text=label_text,
            bg=DUPLICATE_DIALOG_BG,
            fg=field_name_color
        ).grid(row=index, column=0, sticky="e", padx=10, pady=2)

        tk.Label(
            dialog,
            text=str(existing_value),
            bg=DUPLICATE_DIALOG_BG,
            fg=FG_COLOR,
            anchor="w",
            justify="left",
            wraplength=210
        ).grid(row=index, column=1, sticky="w", padx=10, pady=2)

        if existing_compare != csv_compare:
            csv_button = ColorButton(
                dialog,
                text="CSV",
                width=4,
                bg=CSV_BUTTON_NORMAL_BG,
                fg=CSV_BUTTON_NORMAL_FG,
                activebackground=EXPORT_ACTIVE_BG,
                activeforeground=EXPORT_FG,
                command=lambda f=field_name: toggle_csv_field(f)
            )

            csv_button.grid(
                row=index,
                column=2,
                sticky="",
                padx=10,
                pady=2
            )

            csv_buttons[field_name] = csv_button
                    
        tk.Label(
            dialog,
            text=str(csv_value),
            bg=DUPLICATE_DIALOG_BG,
            fg="aqua",
            anchor="w",
            justify="left",
            wraplength=210
        ).grid(row=index, column=3, sticky="w", padx=10, pady=2)

    def replace_all():
        if any(var.get() == 1 for var in selected_vars.values()):
            return

        result["action"] = "replace"
        dialog.destroy()
        
    def skip_record():
        result["action"] = "skip"
        dialog.destroy()

    def add_as_new_record():
        result["action"] = "add_new"
        dialog.destroy()

    def update_selected():
        selected = [
            field_name
            for field_name, var in selected_vars.items()
            if var.get() == 1
        ]

        if not selected:
            show_error(
                "No Fields Selected",
                "Select at least one CSV field to update, or choose Skip."
            )
            return

        result["action"] = "selected"
        result["selected_fields"] = selected
        dialog.destroy()

    def cancel_import():
        result["action"] = "cancel"
        dialog.destroy()

    action_frame = tk.Frame(
        dialog,
        bg=DUPLICATE_DIALOG_BG
    )

    action_frame.grid(
        row=1,
        column=4,
        rowspan=len(field_labels) + 2,
        sticky="n",
        padx=(18, 10),
        pady=(0, 10)
    )

    tk.Label(
        action_frame,
        text="Actions",
        bg=DUPLICATE_DIALOG_BG,
        fg=FG_COLOR
    ).grid(
        row=0,
        column=0,
        pady=(0, 12)
    )

    replace_all_button = ColorButton(
        action_frame,
        text="Replace Entire\nExisting Record",
        bg=EXPORT_BG,
        fg=EXPORT_FG,
        activebackground=EXPORT_ACTIVE_BG,
        activeforeground=EXPORT_FG,
        width=22,
        command=replace_all
    )

    replace_all_button.grid(
        row=1,
        column=0,
        sticky="ew",
        pady=(0, 10)
    )
    
    update_selected_button = ColorButton(
        action_frame,
        text="Update Selected\nFields Only",
        bg=EXPORT_BG,
        fg=EXPORT_FG,
        activebackground=EXPORT_ACTIVE_BG,
        activeforeground=EXPORT_FG,
        width=22,
        command=update_selected,
        state="disabled"
    )

    update_selected_button.grid(
        row=2,
        column=0,
        sticky="ew",
        pady=(0, 10)
    )
    
    update_action_button_states()    
    
    ColorButton(
        action_frame,
        text="Add as Separate\nNew Record",
        bg=SAVE_BG,
        fg=SAVE_FG,
        activebackground=SAVE_ACTIVE_BG,
        activeforeground=SAVE_ACTIVE_FG,
        width=22,
        command=add_as_new_record
    ).grid(
        row=3,
        column=0,
        sticky="ew",
        pady=(0, 10)
    )

    ColorButton(
        action_frame,
        text="Skip CSV Duplicate",
        bg=SHOW_ALL_RECORDS_BG,
        fg=SHOW_ALL_RECORDS_FG,
        activebackground=SHOW_ALL_RECORDS_ACTIVE_BG,
        activeforeground=SHOW_ALL_RECORDS_FG,
        width=22,
        command=skip_record
    ).grid(
        row=4,
        column=0,
        sticky="ew",
        pady=(0, 10)
    )

    ColorButton(
        action_frame,
        text="Cancel Import",
        bg=RESET_BG,
        fg=RESET_FG,
        activebackground=RESET_ACTIVE_BG,
        activeforeground=RESET_FG,
        width=22,
        command=cancel_import
    ).grid(
        row=5,
        column=0,
        sticky="ew"
    )
    
    center_toplevel_over_parent(dialog, root)
    dialog.grab_set()
    dialog.wait_window()
    
    return result

def format_record_for_report(record):
    field_order = [
    ("honorific_title", "Title"),
    ("first_name", "First Name"),
    ("last_name", "Last Name"),
    ("area_code", "Area Code"),
    ("phone", "Phone"),
    ("email", "Email"),
    ("mailaddress_line1", "Address 1"),
    ("mailaddress_line2", "Address 2"),
    ("city", "City"),
    ("state", "State"),
    ("zip_code", "ZIP Code"),
    ("email_list", "Email List"),
    ("mail_list", "Mail List"),
    ("notes", "Notes")
]

    lines = []

    for field_name, label in field_order:
        value = record.get(field_name, "")

        if field_name == "area_code":
            value = format_area_code(value)

        if field_name == "phone":
            value = format_phone(value)

        lines.append(f"{label}: {value}")

    return "\n".join(lines)

def write_csv_update_exceptions_report(
    source_csv_path,
    not_added_records,
    partially_added_records
):
    report_timestamp = datetime.now().strftime("%Y-%m-%d %H%M")
    default_filename = f"CSV Update Exceptions {report_timestamp}.txt"

    report_path = filedialog.asksaveasfilename(
        parent=root,
        defaultextension=".txt",
        initialfile=default_filename,
        filetypes=[("Text files", "*.txt")]
    )

    if not report_path:
        return None

    with open(report_path, "w", encoding="utf-8") as report:
        report.write("CSV UPDATE EXCEPTIONS REPORT\n")
        report.write("=" * 80 + "\n\n")
        report.write(f"Source CSV file:\n{source_csv_path}\n\n")
        report.write(f"Report created:\n{datetime.now().strftime('%c')}\n\n")

        report.write("CSV RECORDS NOT ADDED TO DATABASE\n")
        report.write("=" * 80 + "\n\n")

        if not not_added_records:
            report.write("None.\n\n")
        else:
            for index, item in enumerate(not_added_records, start=1):
                report.write(f"NOT ADDED RECORD {index}\n")
                report.write("-" * 80 + "\n")
                report.write(f"Reason: {item['reason']}\n\n")
                report.write(format_record_for_report(item["csv_record"]))
                report.write("\n\n")

        report.write("\n")
        report.write("CSV RECORDS PARTIALLY ADDED TO DATABASE\n")
        report.write("=" * 80 + "\n\n")

        if not partially_added_records:
            report.write("None.\n\n")
        else:
            for index, item in enumerate(partially_added_records, start=1):
                report.write(f"PARTIALLY ADDED RECORD {index}\n")
                report.write("-" * 80 + "\n")
                report.write(f"Existing database record ID: {item['existing_id']}\n\n")

                report.write("Accepted / Added from CSV:\n")
                report.write("-" * 40 + "\n")

                if item["accepted_fields"]:
                    for field_name in item["accepted_fields"]:
                        label = item["field_labels"].get(field_name, field_name)
                        value = item["csv_record"].get(field_name, "")

                        if field_name == "area_code":
                            value = format_area_code(value)

                        if field_name == "phone":
                            value = format_phone(value)

                        report.write(f"{label}: {value}\n")
                else:
                    report.write("None.\n")

                report.write("\nNot selected / Not added from CSV:\n")
                report.write("-" * 40 + "\n")

                if item["rejected_fields"]:
                    for field_name in item["rejected_fields"]:
                        label = item["field_labels"].get(field_name, field_name)
                        value = item["csv_record"].get(field_name, "")

                        if field_name == "area_code":
                            value = format_area_code(value)

                        if field_name == "phone":
                            value = format_phone(value)

                        report.write(f"{label}: {value}\n")
                else:
                    report.write("None.\n")

                report.write("\nFull CSV record:\n")
                report.write("-" * 40 + "\n")
                report.write(format_record_for_report(item["csv_record"]))
                report.write("\n\n")

    open_file(report_path)
    return report_path

def import_new_records_from_csv():
    file_path = filedialog.askopenfilename(
        parent=root,
        filetypes=[("CSV files", "*.csv")]
    )

    if not file_path:
        return

    new_count = 0
    skipped_count = 0
    matching_skipped_count = 0
    replaced_count = 0
    selected_update_count = 0
    no_list_rejected_count = 0

    not_added_records = []
    partially_added_records = []

    field_labels = {
        "honorific_title": "Title",
        "first_name": "First Name",
        "last_name": "Last Name",
        "area_code": "Area Code",
        "phone": "Phone",
        "email": "Email",
        "mailaddress_line1": "Address 1",
        "mailaddress_line2": "Address 2",
        "city": "City",
        "state": "State",
        "zip_code": "ZIP Code",
        "email_list": "Email List",
        "mail_list": "Mail List",
        "notes": "Notes"
    }

    with open(file_path, "r", newline="", encoding="utf-8-sig") as csvfile:
        reader = csv.DictReader(csvfile)

        with sqlite3.connect(DB_FILE) as conn:
            for row in reader:

                if not any(str(value).strip() for value in row.values()):
                    continue

                csv_record = {
                    "honorific_title": row.get("honorific_title", "").strip(),
                    "first_name": row.get("first_name", "").strip(),
                    "last_name": row.get("last_name", "").strip(),
                    "area_code": row.get("area_code", "").strip(),
                    "phone": row.get("phone", "").strip(),
                    "email": row.get("email", "").strip(),
                    "mailaddress_line1": row.get("mailaddress_line1", "").strip(),
                    "mailaddress_line2": row.get("mailaddress_line2", "").strip(),
                    "city": row.get("city", "").strip(),
                    "state": row.get("state", "").strip(),
                    "zip_code": row.get("zip_code", "").strip(),
                    "email_list": csv_int(row.get("email_list", 0)),
                    "mail_list": csv_int(row.get("mail_list", 0)),
                    "notes": row.get("notes", "").strip()
                }

                cursor = conn.execute("""
                    SELECT id, honorific_title, first_name, last_name,
                        area_code, phone, email, mailaddress_line1,
                        mailaddress_line2, city, state, zip_code, email_list, mail_list, notes
                    FROM recipients
                    WHERE (
                            lower(email) = lower(?)
                            AND email != ''
                        )
                    OR (
                            phone = ?
                            AND phone != ''
                        )
                    OR (
                            lower(last_name) = lower(?)
                            AND last_name != ''
                            AND (
                                (
                                    ? != ''
                                    AND (
                                        lower(mailaddress_line1) = lower(?)
                                        OR lower(mailaddress_line2) = lower(?)
                                    )
                                )
                                OR
                                (
                                    ? != ''
                                    AND (
                                        lower(mailaddress_line1) = lower(?)
                                        OR lower(mailaddress_line2) = lower(?)
                                    )
                                )
                            )
                        )
                """, (
                    csv_record["email"],
                    csv_record["phone"],
                    csv_record["last_name"],

                    csv_record["mailaddress_line1"],
                    csv_record["mailaddress_line1"],
                    csv_record["mailaddress_line1"],

                    csv_record["mailaddress_line2"],
                    csv_record["mailaddress_line2"],
                    csv_record["mailaddress_line2"]
                ))

                duplicate_rows = cursor.fetchall()

                if duplicate_rows:
                    duplicate_candidates = [
                        database_duplicate_row_to_record(duplicate)
                        for duplicate in duplicate_rows
                    ]

                    exact_match = next(
                        (
                            (existing_id, existing_record)
                            for existing_id, existing_record in duplicate_candidates
                            if csv_record_matches_existing_record(
                                existing_record,
                                csv_record
                            )
                        ),
                        None
                    )

                    if exact_match:
                        skipped_count += 1
                        matching_skipped_count += 1

                        not_added_records.append({
                            "reason": "Matching Records Skipped Without User Review",
                            "csv_record": csv_record.copy()
                        })

                        continue

                    existing_id, existing_record = max(
                        duplicate_candidates,
                        key=lambda candidate: (
                            csv_duplicate_match_score(
                                candidate[1],
                                csv_record
                            ),
                            -candidate[0]
                        )
                    )

                    decision = choose_csv_fields_for_update(
                        existing_record,
                        csv_record
                    )

                    if decision["action"] == "cancel":
                        not_added_records.append({
                            "reason": "Import canceled by user while reviewing this duplicate.",
                            "csv_record": csv_record.copy()
                        })
                        break

                    if decision["action"] == "skip":
                        skipped_count += 1

                        not_added_records.append({
                            "reason": "Skipped by user after duplicate comparison.",
                            "csv_record": csv_record.copy()
                        })

                        continue

                    if decision["action"] == "add_new":
                        if not csv_record["email_list"] and not csv_record["mail_list"]:
                            show_error(
                                "Invalid New Record",
                                "Adding this CSV record would create a record on no list.\n\n"
                                "The new record was skipped."
                            )

                            skipped_count += 1
                            no_list_rejected_count += 1

                            not_added_records.append({
                                "reason": "New record skipped because CSV record had neither Email List nor Mail List selected.",
                                "csv_record": csv_record.copy()
                            })

                            continue

                        insert_csv_record(conn, csv_record)
                        new_count += 1
                        continue

                    if decision["action"] == "replace":
                        if not csv_record["email_list"] and not csv_record["mail_list"]:
                            show_error(
                                "Invalid Replace",
                                "Replacing with this CSV record would leave the record on no list.\n\n"
                                "The replacement was skipped."
                            )

                            skipped_count += 1
                            no_list_rejected_count += 1

                            not_added_records.append({
                                "reason": "Replacement skipped because CSV record had neither Email List nor Mail List selected.",
                                "csv_record": csv_record.copy()
                            })

                            continue

                        conn.execute("""
                            UPDATE recipients
                            SET honorific_title = ?,
                                first_name = ?,
                                last_name = ?,
                                area_code = ?,
                                phone = ?,
                                email = ?,
                                mailaddress_line1 = ?,
                                mailaddress_line2 = ?,
                                city = ?,
                                state = ?,
                                zip_code = ?,
                                email_list = ?,
                                mail_list = ?,
                                notes = ?
                            WHERE id = ?
                        """, (
                            csv_record["honorific_title"],
                            csv_record["first_name"],
                            csv_record["last_name"],
                            csv_record["area_code"],
                            csv_record["phone"],
                            csv_record["email"],
                            csv_record["mailaddress_line1"],
                            csv_record["mailaddress_line2"],
                            csv_record["city"],
                            csv_record["state"],
                            csv_record["zip_code"],
                            csv_record["email_list"],
                            csv_record["mail_list"],
                            csv_record["notes"],
                            existing_id
                        ))

                        replaced_count += 1
                        continue

                    if decision["action"] == "selected":
                        updated_record = existing_record.copy()

                        for field_name in decision["selected_fields"]:
                            updated_record[field_name] = csv_record[field_name]

                        if (
                            not updated_record["email_list"]
                            and not updated_record["mail_list"]
                        ):
                            show_error(
                                "Invalid Update",
                                "The selected update would leave the record on no list.\n\n"
                                "The update was skipped."
                            )

                            skipped_count += 1
                            no_list_rejected_count += 1

                            not_added_records.append({
                                "reason": "Selective update skipped because it would leave the record on no list.",
                                "csv_record": csv_record.copy()
                            })

                            continue

                        conn.execute("""
                            UPDATE recipients
                            SET honorific_title = ?,
                                first_name = ?,
                                last_name = ?,
                                area_code = ?,
                                phone = ?,
                                email = ?,
                                mailaddress_line1 = ?,
                                mailaddress_line2 = ?,
                                city = ?,
                                state = ?,
                                zip_code = ?,
                                email_list = ?,
                                mail_list = ?,
                                notes = ?
                            WHERE id = ?
                        """, (
                            updated_record["honorific_title"],
                            updated_record["first_name"],
                            updated_record["last_name"],
                            updated_record["area_code"],
                            updated_record["phone"],
                            updated_record["email"],
                            updated_record["mailaddress_line1"],
                            updated_record["mailaddress_line2"],
                            updated_record["city"],
                            updated_record["state"],
                            updated_record["zip_code"],
                            updated_record["email_list"],
                            updated_record["mail_list"],
                            updated_record["notes"],
                            existing_id
                        ))

                        accepted_fields = decision["selected_fields"]

                        all_fields = [
                            "honorific_title",
                            "first_name",
                            "last_name",
                            "area_code",
                            "phone",
                            "email",
                            "mailaddress_line1",
                            "mailaddress_line2",
                            "city",
                            "state",
                            "zip_code",
                            "email_list",
                            "mail_list",
                            "notes"
                        ]

                        rejected_fields = [
                            field_name
                            for field_name in all_fields
                            if field_name not in accepted_fields
                        ]

                        partially_added_records.append({
                            "existing_id": existing_id,
                            "csv_record": csv_record.copy(),
                            "accepted_fields": accepted_fields,
                            "rejected_fields": rejected_fields,
                            "field_labels": field_labels
                        })

                        selected_update_count += 1
                        continue

                else:
                    if not csv_record["email_list"] and not csv_record["mail_list"]:
                        no_list_rejected_count += 1
                        skipped_count += 1

                        not_added_records.append({
                            "reason": "Rejected because neither Email List nor Mail List was selected.",
                            "csv_record": csv_record.copy()
                        })

                        continue

                    insert_csv_record(conn, csv_record)
                    new_count += 1

        changes_made = (
            new_count > 0
            or replaced_count > 0
            or selected_update_count > 0
        )

        if changes_made:
            mark_unbacked_changes()

        refresh_database_status(update_timestamp=changes_made)

    exceptions_report_path = write_csv_update_exceptions_report(
        file_path,
        not_added_records,
        partially_added_records
    )

    report_message = ""

    if exceptions_report_path:
        report_message = f"\n\nExceptions report saved:\n{exceptions_report_path}"

    show_info(
        "Import Complete",
        f"New records added: {new_count}\n"
        f"Existing records replaced: {replaced_count}\n"
        f"Existing records selectively updated: {selected_update_count}\n"
        f"Matching records skipped without user review: {matching_skipped_count}\n"
        f"Total records skipped: {skipped_count}\n"
        f"Rejected for no list selection: {no_list_rejected_count}"
        f"{report_message}"
    )

def delete_record():
    global CURRENT_RECORD_ID

    record_id = CURRENT_RECORD_ID or edit_id_entry.get().strip()

    if not str(record_id).isdigit():
        show_error(
            "Error",
            "No Record is currently loaded.\n\nPlease enter a numeric record ID."
        )
        return

    CURRENT_RECORD_ID = record_id
    
    confirm = ask_yes_no(
        "Confirm Delete",
        f"Are you sure you want to delete record ID {record_id}?\n\nTHIS CANNOT BE UNDONE!"
    )

    if not confirm:
        return

    with sqlite3.connect(DB_FILE) as conn:
        cursor = conn.execute(
            "DELETE FROM recipients WHERE id = ?",
            (record_id,)
        )

    if cursor.rowcount == 0:
        show_error(
            "Error",
            f"No record found with ID {record_id}."
        )
    else:
        edit_id_entry.delete(0, tk.END)
        clear_form()
        mark_unbacked_changes()
        refresh_database_status(update_timestamp=True)
        refresh_current_search_results()
        save_button.config(state="normal")
        delete_button.config(state="disabled")

        edit_id_entry.focus_set()

        show_info(
            "Deleted",
            f"Record ID {record_id} deleted."
        )

def load_record_for_edit():
    global CURRENT_RECORD_ID

    record_id = edit_id_entry.get().strip()
    
    if not record_id.isdigit():
        show_error(
            "Error",
            "Invalid ID.\n\nPlease enter a numeric record ID."
        )
        return

    with sqlite3.connect(DB_FILE) as conn:
        cursor = conn.execute("""
            SELECT honorific_title, first_name, last_name, area_code, phone,
                   email, mailaddress_line1, mailaddress_line2, city, state, zip_code,
                   email_list, mail_list, notes
            FROM recipients
            WHERE id = ?
        """, (record_id,))

        record = cursor.fetchone()

    if not record:
        show_error(
            "Error",
            f"No record found with ID {record_id}."
        )
        return

    clear_form()

    edit_id_entry.delete(0, tk.END)
    edit_id_entry.insert(0, record_id)

    CURRENT_RECORD_ID = record_id
    
    (
        hon_title, first, last, areacode, phone, email,
        address1, address2, city, state, zip_code, listemail, listmail, notes
    ) = record

    honorific_title_entry.insert(0, hon_title or "")
    first_name_entry.insert(0, first or "")
    last_name_entry.insert(0, last or "")
    area_code_entry.insert(0, areacode or "")
    phone_entry.insert(0, phone or "")
    email_entry.insert(0, email or "")
    mailaddress_line1_entry.insert(0, address1 or "")
    mailaddress_line2_entry.insert(0, address2 or "")
    city_entry.insert(0, city or "")
    state_entry.insert(0, state or "")
    zip_code_entry.insert(0, zip_code or "")

    email_list_var.set(listemail)
    mail_list_var.set(listmail)
    notes_text.insert("1.0", notes or "")
    
    save_button.config(text="Update Record")
    save_button.config(state="normal")
    delete_button.config(state="normal")

def update_record():
    global CURRENT_RECORD_ID

    record_id = CURRENT_RECORD_ID or edit_id_entry.get().strip()

    if not str(record_id).isdigit():
        show_error(
            "Error",
            "No record is currently loaded.\n\nLoad a record before updating."
        )
        return

    CURRENT_RECORD_ID = record_id
    
    form_data = get_form_data_or_error()

    if form_data is None:
        return

    with sqlite3.connect(DB_FILE) as conn:
        cursor = conn.execute("""
            UPDATE recipients
            SET honorific_title = ?,
                first_name = ?,
                last_name = ?,
                area_code = ?,
                phone = ?,
                email = ?,
                mailaddress_line1 = ?,
                mailaddress_line2 = ?,
                city = ?,
                state = ?,
                zip_code = ?,
                email_list = ?,
                mail_list = ?,
                notes = ?
            WHERE id = ?
        """, (*form_data, record_id))

    if cursor.rowcount == 0:
        show_error(
            "Error",
            f"No record found with ID {record_id}."
        )
    else:
        show_info(
            "Updated",
            f"Record ID {record_id} updated."
        )
        mark_unbacked_changes()
        refresh_database_status(update_timestamp=True)
        clear_form()
        refresh_current_search_results()
        edit_id_entry.focus_set()

def save_or_update_record():
    if CURRENT_RECORD_ID:
        update_record()
    else:
        save_recipient()

def add_mail_label_page(
    document,
    recipient_rows,
    start_index,
    table_rows,
    table_cols,
    row_height_inches,
    cell_width_inches
):
    table = document.add_table(rows=table_rows, cols=table_cols)
    table.alignment = 1
    table.autofit = False
    table.allow_autofit = False

    label_index = start_index

    for row in table.rows:
        row.height = Inches(row_height_inches)

        for cell in row.cells:
            cell.width = Inches(cell_width_inches)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            if label_index >= len(recipient_rows):
                continue

            hon_title, first, last, address1, address2, city, state, zip_code = recipient_rows[label_index]

            full_name = " ".join(
                part for part in [hon_title, first, last]
                if part
            )

            lines = [
                full_name,
                address1,
                address2,
                f"{city}, {state}  {zip_code}"
            ]

            lines = [line for line in lines if line]

            longest_line = max(lines, key=len)
            average_char_width_inches = 0.075
            longest_line_width_inches = len(longest_line) * average_char_width_inches

            left_indent_inches = (
                cell_width_inches - longest_line_width_inches
            ) / 2

            left_indent_inches = max(left_indent_inches, 0)

            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.left_indent = Inches(left_indent_inches)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

            for line_index, line in enumerate(lines):
                if line_index < len(lines) - 1:
                    paragraph.add_run(line + "\n")
                else:
                    paragraph.add_run(line)

            label_index += 1

    return label_index


def export_mail_labels_5162():
    file_path = filedialog.asksaveasfilename(
        parent=root,
        defaultextension=".docx",
        initialfile=f"Newsletter 14 per page {datetime.now().strftime('%Y-%m-%d %H%M')}.docx",
        filetypes=[("Word documents", "*.docx")]
    )

    if not file_path:
        return

    with sqlite3.connect(DB_FILE) as conn:
        cursor = conn.execute("""
            SELECT honorific_title, first_name, last_name,
                mailaddress_line1, mailaddress_line2, city, state, zip_code
            FROM recipients
            WHERE mail_list = 1
            ORDER BY lower(last_name), lower(first_name)
        """)
        rows = cursor.fetchall()

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.83)
    section.bottom_margin = Inches(0.57)
    section.left_margin = Inches(0.16)
    section.right_margin = Inches(0.16)

    label_index = 0

    while label_index < len(rows):
        if label_index > 0:
            document.add_page_break()

        label_index = add_mail_label_page(
            document=document,
            recipient_rows=rows,
            start_index=label_index,
            table_rows=7,
            table_cols=2,
            row_height_inches=1.333,
            cell_width_inches=4.0
        )

    document.save(file_path)
    open_file(file_path)


def export_mail_labels_5160():
    file_path = filedialog.asksaveasfilename(
        parent=root,
        defaultextension=".docx",
        initialfile=f"Newsletter 30 per page {datetime.now().strftime('%Y-%m-%d %H%M')}.docx",
        filetypes=[("Word documents", "*.docx")]
    )

    if not file_path:
        return

    with sqlite3.connect(DB_FILE) as conn:
        cursor = conn.execute("""
            SELECT honorific_title, first_name, last_name,
                   mailaddress_line1, mailaddress_line2, city, state, zip_code
            FROM recipients
            WHERE mail_list = 1
            ORDER BY lower(last_name), lower(first_name)
        """)
        rows = cursor.fetchall()

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.30)
    section.left_margin = Inches(0.30)
    section.right_margin = Inches(0.30)

    label_index = 0

    while label_index < len(rows):
        if label_index > 0:
            document.add_page_break()

        label_index = add_mail_label_page(
            document=document,
            recipient_rows=rows,
            start_index=label_index,
            table_rows=10,
            table_cols=3,
            row_height_inches=0.99,
            cell_width_inches=2.625
        )

    document.save(file_path)
    open_file(file_path)

setup_database()

root = tk.Tk()

root.title(f"{APP_NAME} {APP_VERSION}")
root.configure(bg=BG_COLOR)

def get_active_app_window():
    """
    Return the currently active window within this application.

    If a custom dialog is open, message boxes should appear above that
    dialog. Otherwise, they should appear above the main application window.
    """
    try:
        focused_widget = root.focus_get()

        if focused_widget is not None:
            return focused_widget.winfo_toplevel()

    except tk.TclError:
        pass

    return root


def center_toplevel_over_parent(dialog, parent=None):
    """
    Center a custom Toplevel window over its parent window.

    This uses the parent's current screen coordinates, so it also works
    when the main application window has been moved to another display.
    """
    if parent is None:
        parent = root

    parent.update_idletasks()
    dialog.update_idletasks()

    parent_x = parent.winfo_rootx()
    parent_y = parent.winfo_rooty()
    parent_width = parent.winfo_width()
    parent_height = parent.winfo_height()

    dialog_width = dialog.winfo_reqwidth()
    dialog_height = dialog.winfo_reqheight()

    x_position = parent_x + int((parent_width - dialog_width) / 2)
    y_position = parent_y + int((parent_height - dialog_height) / 2)

    # The signed format is important for monitors positioned to the
    # left of or above the primary display, where coordinates are negative.
    dialog.geometry(f"{x_position:+d}{y_position:+d}")

    dialog.deiconify()
    dialog.lift()
    dialog.focus_force()


def show_error(title, message, parent=None):
    if parent is None:
        parent = get_active_app_window()

    return messagebox.showerror(
        title,
        message,
        parent=parent
    )


def show_warning(title, message, parent=None):
    if parent is None:
        parent = get_active_app_window()

    return messagebox.showwarning(
        title,
        message,
        parent=parent
    )


def show_info(title, message, parent=None):
    if parent is None:
        parent = get_active_app_window()

    return messagebox.showinfo(
        title,
        message,
        parent=parent
    )


def ask_yes_no(title, message, parent=None):
    if parent is None:
        parent = get_active_app_window()

    return messagebox.askyesno(
        title,
        message,
        parent=parent
    )
    

def insert_keypad_digit(event, digit):
    """
    Insert a digit from the extended numeric keypad.

    On Windows, keypad keys may be reported differently depending
    on whether Num Lock is on or off. Returning "break" prevents
    Tk from also processing the same keystroke as navigation.
    """
    try:
        event.widget.insert(tk.INSERT, digit)
        return "break"
    except tk.TclError:
        return None


def enable_windows_numeric_keypad():
    if os.name != "nt":
        return

    keypad_digit_map = {
        # Num Lock on
        "KP_0": "0",
        "KP_1": "1",
        "KP_2": "2",
        "KP_3": "3",
        "KP_4": "4",
        "KP_5": "5",
        "KP_6": "6",
        "KP_7": "7",
        "KP_8": "8",
        "KP_9": "9",

        # Num Lock off: Windows may report navigation-style keysyms
        "KP_Insert": "0",
        "KP_End": "1",
        "KP_Down": "2",
        "KP_Next": "3",
        "KP_Left": "4",
        "KP_Begin": "5",
        "KP_Right": "6",
        "KP_Home": "7",
        "KP_Up": "8",
        "KP_Prior": "9",
    }

    for widget_class in ("Entry", "TEntry", "Text"):
        for keysym, digit in keypad_digit_map.items():
            root.bind_class(
                widget_class,
                f"<KeyPress-{keysym}>",
                lambda event, value=digit: insert_keypad_digit(event, value)
            )


enable_windows_numeric_keypad()


def show_about_dialog():
    root.lift()
    root.update_idletasks()

    show_info(
        f"About {APP_NAME}",
        f"{APP_NAME}\n"
        f"Version {APP_VERSION}\n\n"
        f"Copyright © 2026 William Tinney\n\n"
        f"Licensed under the MIT License.\n"
        f"SPDX-License-Identifier: MIT\n\n"
        f"Support: william.tinney@gmail.com",
        parent=root
    )
    
menu_bar = tk.Menu(root)

help_menu = tk.Menu(menu_bar, tearoff=0)

help_menu.add_command(
    label="User Manual",
    command=open_user_manual
)

help_menu.add_separator()

help_menu.add_command(
    label=f"About {APP_NAME}",
    command=show_about_dialog
)

menu_bar.add_cascade(
    label="Help",
    menu=help_menu
)

root.config(menu=menu_bar)

    
def choose_fixed_font(size=12):
    available_fonts = set(tkfont.families())

    preferred_fonts = [
        "Menlo",          # macOS
        "Consolas",       # Windows
        "Courier New",    # macOS / Windows fallback
        "TkFixedFont"
    ]

    for font_name in preferred_fonts:
        if font_name in available_fonts:
            return (font_name, size)

    return ("TkFixedFont", size)


SEARCH_RESULTS_FONT = choose_fixed_font(9)
SEARCH_RESULTS_HEADER_FONT = choose_fixed_font(9)

# Optional: leave enough room for the wide layout
root.columnconfigure(0, weight=0)
root.columnconfigure(1, weight=0)
root.columnconfigure(2, weight=1)
root.columnconfigure(3, weight=0)

root.rowconfigure(4, weight=1)


# -------------------------
# LEFT FORM AREA
# -------------------------

form_frame = tk.Frame(root, bg=BG_COLOR)
form_frame.grid(row=0, column=0, sticky="nw", padx=20, pady=20)

preserve_title_var = tk.IntVar()
preserve_first_name_var = tk.IntVar()
preserve_last_name_var = tk.IntVar()
preserve_address1_var = tk.IntVar()
preserve_address2_var = tk.IntVar()
preserve_city_var = tk.IntVar()

tk.Label(form_frame, text="Title", bg=BG_COLOR, fg=FG_COLOR).grid(
    row=0, column=0, sticky="w", pady=4
)
honorific_title_entry = tk.Entry(form_frame, width=22)
honorific_title_entry.grid(row=0, column=1, sticky="w", padx=10, pady=4)

tk.Checkbutton(
    form_frame,
    text="Exact",
    bg=CHECKBOX_BG,
    fg=CHECKBOX_FG,
    activebackground=BG_COLOR,
    selectcolor=BG_COLOR,
    variable=preserve_title_var,
    takefocus=0,
).grid(row=0, column=2, sticky="w", padx=5, pady=4)

tk.Label(form_frame, text="First name", bg=BG_COLOR, fg=FG_COLOR).grid(
    row=1, column=0, sticky="w", pady=4
)
first_name_entry = tk.Entry(form_frame, width=22)
first_name_entry.grid(row=1, column=1, sticky="w", padx=10, pady=4)

tk.Checkbutton(
    form_frame,
    text="Exact",
    bg=CHECKBOX_BG,
    fg=CHECKBOX_FG,
    activebackground=BG_COLOR,
    selectcolor=BG_COLOR,
    variable=preserve_first_name_var,
    takefocus=0,
).grid(row=1, column=2, sticky="w", padx=5, pady=4)

tk.Label(form_frame, text="Last name", bg=BG_COLOR, fg=FG_COLOR).grid(
    row=2, column=0, sticky="w", pady=4
)
last_name_entry = tk.Entry(form_frame, width=22)
last_name_entry.grid(row=2, column=1, sticky="w", padx=10, pady=4)

tk.Checkbutton(
    form_frame,
    text="Exact",
    bg=CHECKBOX_BG,
    fg=CHECKBOX_FG,
    activebackground=BG_COLOR,
    selectcolor=BG_COLOR,
    variable=preserve_last_name_var,
    takefocus=0,
).grid(row=2, column=2, sticky="w", padx=5, pady=4)

tk.Label(form_frame, text="Area code", bg=BG_COLOR, fg=FG_COLOR).grid(
    row=3, column=0, sticky="w", pady=4
)
area_code_entry = tk.Entry(form_frame, width=10)
area_code_entry.grid(row=3, column=1, sticky="w", padx=10, pady=4)

tk.Label(form_frame, text="Phone number", bg=BG_COLOR, fg=FG_COLOR).grid(
    row=4, column=0, sticky="w", pady=4
)
phone_entry = tk.Entry(form_frame, width=13)
phone_entry.grid(row=4, column=1, sticky="w", padx=10, pady=4)

tk.Label(form_frame, text="Email Address", bg=BG_COLOR, fg=FG_COLOR).grid(
    row=5, column=0, sticky="w", pady=4
)
email_entry = tk.Entry(form_frame, width=34)
email_entry.grid(row=5, column=1, sticky="w", padx=10, pady=4)

tk.Label(form_frame, text="Address 1", bg=BG_COLOR, fg=FG_COLOR).grid(
    row=6, column=0, sticky="w", pady=4
)
mailaddress_line1_entry = tk.Entry(form_frame, width=34)
mailaddress_line1_entry.grid(row=6, column=1, sticky="w", padx=10, pady=4)

tk.Checkbutton(
    form_frame,
    text="Exact",
    bg=CHECKBOX_BG,
    fg=CHECKBOX_FG,
    activebackground=BG_COLOR,
    selectcolor=BG_COLOR,
    variable=preserve_address1_var,
    takefocus=0,
).grid(row=6, column=2, sticky="w", padx=5, pady=4)

tk.Label(form_frame, text="Address 2 (Opt)", bg=BG_COLOR, fg=FG_COLOR).grid(
    row=7, column=0, sticky="w", pady=4
)
mailaddress_line2_entry = tk.Entry(form_frame, width=34)
mailaddress_line2_entry.grid(row=7, column=1, sticky="w", padx=10, pady=4)

tk.Checkbutton(
    form_frame,
    text="Exact",
    bg=CHECKBOX_BG,
    fg=CHECKBOX_FG,
    activebackground=BG_COLOR,
    selectcolor=BG_COLOR,
    variable=preserve_address2_var,
    takefocus=0,
).grid(row=7, column=2, sticky="w", padx=5, pady=4)

tk.Label(form_frame, text="City", bg=BG_COLOR, fg=FG_COLOR).grid(
    row=8, column=0, sticky="w", pady=4
)
city_entry = tk.Entry(form_frame, width=22)
city_entry.grid(row=8, column=1, sticky="w", padx=10, pady=4)

tk.Checkbutton(
    form_frame,
    text="Exact",
    bg=CHECKBOX_BG,
    fg=CHECKBOX_FG,
    activebackground=BG_COLOR,
    selectcolor=BG_COLOR,
    variable=preserve_city_var,
    takefocus=0,
).grid(row=8, column=2, sticky="w", padx=5, pady=4)

state_var = tk.StringVar()

def force_uppercase(*args):
    state_var.set(state_var.get().upper())

state_var.trace_add("write", force_uppercase)

tk.Label(form_frame, text="State", bg=BG_COLOR, fg=FG_COLOR).grid(
    row=9, column=0, sticky="w", pady=4
)
state_entry = tk.Entry(form_frame, width=5, textvariable=state_var)
state_entry.grid(row=9, column=1, sticky="w", padx=10, pady=4)

tk.Label(form_frame, text="5 digit ZIP code", bg=BG_COLOR, fg=FG_COLOR).grid(
    row=10, column=0, sticky="w", pady=4
)
zip_code_entry = tk.Entry(form_frame, width=10)
zip_code_entry.grid(row=10, column=1, sticky="w", padx=10, pady=4)

email_list_var = tk.IntVar()
mail_list_var = tk.IntVar()

checkbox_email = tk.Checkbutton(
    form_frame,
    text="Email list",
    bg=CHECKBOX_BG,
    fg=CHECKBOX_FG,
    activebackground=BG_COLOR,
    selectcolor=BG_COLOR,
    variable=email_list_var
)
checkbox_email.grid(row=11, column=1, sticky="w", padx=10, pady=(14, 4))

checkbox_mail = tk.Checkbutton(
    form_frame,
    text="Mail list",
    bg=CHECKBOX_BG,
    fg=CHECKBOX_FG,
    activebackground=BG_COLOR,
    selectcolor=BG_COLOR,
    variable=mail_list_var
)
checkbox_mail.grid(row=12, column=1, sticky="w", padx=10, pady=4)

tk.Label(form_frame, text="Notes", bg=BG_COLOR, fg=FG_COLOR).grid(
    row=13, column=0, sticky="nw", pady=(8, 4)
)

notes_frame = tk.Frame(form_frame, bg=BG_COLOR)
notes_frame.grid(
    row=13,
    column=1,
    columnspan=2,
    sticky="ew",
    padx=10,
    pady=(8, 4)
)

notes_text = tk.Text(
    notes_frame,
    width=32,
    height=5,
    wrap="word"
)
notes_scrollbar = tk.Scrollbar(
    notes_frame,
    orient="vertical",
    command=notes_text.yview
)
notes_text.configure(yscrollcommand=notes_scrollbar.set)

notes_text.grid(row=0, column=0, sticky="ew")
notes_scrollbar.grid(row=0, column=1, sticky="ns")
notes_frame.columnconfigure(0, weight=1)

form_button_frame = tk.Frame(form_frame, bg=BG_COLOR)
form_button_frame.grid(
    row=14,
    column=0,
    columnspan=2,
    sticky="ew",
    pady=(18, 4)
)

form_button_frame.columnconfigure(0, weight=1)
form_button_frame.columnconfigure(1, weight=0)
form_button_frame.columnconfigure(2, weight=0)
form_button_frame.columnconfigure(3, weight=1)

save_button = ColorButton(
    form_button_frame,
    text="Save Recipient",
    command=save_or_update_record
)
style_button(
    save_button,
    bg=SAVE_BG,
    fg="yellow",
    active_bg=SAVE_ACTIVE_BG,
    active_fg="yellow",
    width=18
)
save_button.grid(row=0, column=1, padx=(0, 8), pady=4)

clear_button = ColorButton(
    form_button_frame,
    text="Reset Form",
    command=clear_form
)
style_button(
    clear_button,
    bg=RESET_BG,
    fg="white",
    active_bg=RESET_ACTIVE_BG,
    active_fg="white",
    width=18
)
clear_button.grid(row=0, column=2, padx=(8, 0), pady=4)


# -------------------------
# SEARCH FRAME
# -------------------------

search_frame = tk.LabelFrame(
    root,
    text="Search by Last Name",
    bg=BG_COLOR,
    fg=FG_COLOR,
    padx=10,
    pady=10
)
search_frame.grid(
    row=0,
    column=1,
    columnspan=2,
    rowspan=3,
    sticky="nsew",
    padx=20,
    pady=20
)

search_last_name_var = tk.StringVar()

tk.Label(
    search_frame,
    text="Last name (partial)",
    bg=BG_COLOR,
    fg=FG_COLOR
).grid(row=0, column=0, sticky="w", padx=5, pady=5)

search_last_name_entry = tk.Entry(
    search_frame,
    width=24,
    textvariable=search_last_name_var
)
search_last_name_entry.grid(row=0, column=1, sticky="w", padx=5, pady=5)

search_results_label_var = tk.StringVar(value="Search Results")

tk.Label(
    search_frame,
    textvariable=search_results_label_var,
    bg=BG_COLOR,
    fg=FG_COLOR
).grid(row=1, column=0, columnspan=4, sticky="w", padx=5, pady=(20, 5))

search_columns = (
    "id",
    "first_name",
    "last_name",
    "area_code",
    "phone",
    "email"
)

search_display_headers = {
    "id": "ID",
    "first_name": "First Name",
    "last_name": "Last Name",
    "area_code": "Area Code",
    "phone": "Phone",
    "email": "Email"
}

search_tree_style = ttk.Style(root)

# The native Windows ttk theme may leave the empty portion of a
# Treeview white even when fieldbackground has been configured.
# The clam theme consistently honors the custom Treeview colors.
if os.name == "nt":
    search_tree_style.theme_use("clam")

search_tree_style.configure(    "Search.Treeview",
    background=BG_COLOR,
    foreground=FG_COLOR,
    fieldbackground=BG_COLOR,
    rowheight=26,
    font=SEARCH_RESULTS_FONT
)

search_tree_style.map(
    "Search.Treeview",
    foreground=[
        ("selected", "black")
    ],
    background=[
        ("selected", "#BFD7FF")
    ]
)

search_tree_style.configure(
    "Search.Treeview.Heading",
    font=SEARCH_RESULTS_HEADER_FONT
)

search_results_box = ttk.Treeview(
    search_frame,
    columns=search_columns,
    show="headings",
    height=16,
    style="Search.Treeview",
    selectmode="extended"
)

for column in search_columns:
    search_results_box.heading(
        column,
        text=search_display_headers[column],
        anchor="w"
    )
    search_results_box.column(
        column,
        anchor="w",
        stretch=False,
        width=100
    )

search_results_box.grid(
    row=2,
    column=0,
    columnspan=4,
    sticky="nsew",
    padx=5,
    pady=5
)

search_frame.columnconfigure(2, weight=1)
search_frame.rowconfigure(2, weight=1)

tk.Label(
    search_frame,
    text="Double-click a result row to load that record into the form.",
    bg=BG_COLOR,
    fg="yellow"
).grid(row=3, column=0, columnspan=4, sticky="w", padx=5, pady=(5, 0))

combine_duplicates_button = ColorButton(
    search_frame,
    text="Combine Duplicates",
    command=combine_selected_duplicates,
    bg=UPDATE_RECORD_BG,
    fg=UPDATE_RECORD_FG,
    activebackground=UPDATE_RECORD_ACTIVE_BG,
    activeforeground=UPDATE_RECORD_FG,
    width=20,
    state="disabled"
)

combine_duplicates_button.grid(
    row=4,
    column=3,
    sticky="e",
    padx=5,
    pady=(8, 0)
)

def update_combine_duplicates_button_state(event=None):
    selected_items = list(search_results_box.selection())

    if len(selected_items) > 2:
        search_results_box.selection_set(selected_items[:2])
        selected_items = list(search_results_box.selection())

    if len(selected_items) == 2:
        combine_duplicates_button.config(state="normal")
    else:
        combine_duplicates_button.config(state="disabled")
        
def toggle_second_search_selection(event=None):
    clicked_item = search_results_box.identify_row(event.y)

    if not clicked_item:
        return "break"

    selected_items = list(search_results_box.selection())

    if clicked_item in selected_items:
        search_results_box.selection_remove(clicked_item)
    else:
        if len(selected_items) >= 2:
            search_results_box.selection_set(selected_items[0], clicked_item)
        else:
            search_results_box.selection_add(clicked_item)

    update_combine_duplicates_button_state()
    return "break"

if root.tk.call("tk", "windowingsystem") == "aqua":
    combine_instruction_text = (
        "Select 1st record, then use Cmd+click to select 2nd record to combine."
    )
else:
    combine_instruction_text = (
        "Select 1st record, then use Ctrl+click to select 2nd record to combine."
    )

tk.Label(
    search_frame,
    text=combine_instruction_text,
    bg=BG_COLOR,
    fg="yellow",
    anchor="e",
    justify="right"
).grid(row=4, column=0, columnspan=3, sticky="e", padx=5, pady=(8, 0))

def populate_search_results(rows, headers):
    for item in search_results_box.get_children():
        search_results_box.delete(item)

    if not rows:
        return

    display_rows = []

    for row in rows:
        row = list(row)

        display_rows.append({
            "id": "" if row[0] is None else str(row[0]),
            "first_name": "" if row[1] is None else str(row[1]),
            "last_name": "" if row[2] is None else str(row[2]),
            "area_code": format_area_code("" if row[3] is None else str(row[3])),
            "phone": format_phone("" if row[4] is None else str(row[4])),
            "email": "" if row[5] is None else str(row[5])
        })

    tree_font = tkfont.Font(
        family=SEARCH_RESULTS_FONT[0],
        size=SEARCH_RESULTS_FONT[1]
    )

    for column in search_columns:
        longest_text = search_display_headers[column]

        for row in display_rows:
            if len(row[column]) > len(longest_text):
                longest_text = row[column]

        pixel_width = tree_font.measure(longest_text) + 30

        search_results_box.column(
            column,
            width=pixel_width,
            minwidth=pixel_width,
            stretch=False
        )

    for row in display_rows:
        search_results_box.insert(
            "",
            tk.END,
            values=(
                row["id"],
                row["first_name"],
                row["last_name"],
                row["area_code"],
                row["phone"],
                row["email"]
            )
        )
    
    update_combine_duplicates_button_state()

def search_by_last_name_inline():
    search_term = search_last_name_var.get().strip()

    if not search_term:
        show_error(
            "Search Required",
            "Enter a last name or partial last name."
        )
        return

    with sqlite3.connect(DB_FILE) as conn:
        cursor = conn.execute("""
            SELECT id, first_name, last_name, area_code, phone, email
            FROM recipients
            WHERE lower(last_name) LIKE lower(?)
            ORDER BY lower(last_name), lower(first_name)
        """, (f"%{search_term}%",))

        rows = cursor.fetchall()
        headers = [description[0] for description in cursor.description]

    search_results_label_var.set(f"Search Results ({len(rows)})")
    populate_search_results(rows, headers)

search_last_name_entry.bind(
    "<Return>",
    lambda event: search_by_last_name_inline()
)

search_last_name_entry.bind(
    "<KP_Enter>",
    lambda event: search_by_last_name_inline()
)

def show_all_records_inline():
    with sqlite3.connect(DB_FILE) as conn:
        cursor = conn.execute("""
            SELECT id, first_name, last_name, area_code, phone, email
            FROM recipients
            ORDER BY lower(last_name), lower(first_name)
        """)

        rows = cursor.fetchall()
        headers = [description[0] for description in cursor.description]

    search_results_label_var.set(f"Search Results ({len(rows)})")
    populate_search_results(rows, headers)
    

def refresh_current_search_results():
    search_term = search_last_name_var.get().strip()

    if search_term:
        search_by_last_name_inline()
    else:
        show_all_records_inline()


def load_selected_search_result(event=None):
    selected_item = search_results_box.focus()

    if not selected_item:
        return

    values = search_results_box.item(selected_item, "values")

    if not values:
        return

    record_id = values[0]

    if str(record_id).isdigit():
        edit_id_entry.delete(0, tk.END)
        edit_id_entry.insert(0, record_id)
        load_record_for_edit()


search_results_box.bind("<Double-Button-1>", load_selected_search_result)

search_results_box.bind(
    "<<TreeviewSelect>>",
    update_combine_duplicates_button_state
)

if root.tk.call("tk", "windowingsystem") == "aqua":
    search_results_box.bind(
        "<Command-Button-1>",
        toggle_second_search_selection
    )
else:
    search_results_box.bind(
        "<Control-Button-1>",
        toggle_second_search_selection
    )

search_button = ColorButton(
    search_frame,
    text="Search",
    command=search_by_last_name_inline
)
style_button(
    search_button,
    bg=SHOW_ALL_RECORDS_BG,
    fg=SHOW_ALL_RECORDS_FG,
    active_bg=SHOW_ALL_RECORDS_ACTIVE_BG,
    active_fg=SHOW_ALL_RECORDS_FG,
    width=12
)
search_button.grid(row=0, column=2, sticky="w", padx=8, pady=5)

show_records_button = ColorButton(
    search_frame,
    text="Show All Records",
    command=show_all_records_inline
)
style_button(
    show_records_button,
    bg=SHOW_ALL_RECORDS_BG,
    fg=SHOW_ALL_RECORDS_FG,
    active_bg=SHOW_ALL_RECORDS_ACTIVE_BG,
    active_fg=SHOW_ALL_RECORDS_FG,
    width=16
)
show_records_button.grid(row=0, column=3, sticky="w", padx=8, pady=5)


# -------------------------
# DATABASE STATUS BELOW SEARCH FRAME
# -------------------------

database_path_var = tk.StringVar()
total_records_var = tk.StringVar()
mail_records_var = tk.StringVar()
email_records_var = tk.StringVar()
last_update_var = tk.StringVar()

status_frame = tk.Frame(root, bg=BG_COLOR)
status_frame.grid(row=3, column=1, columnspan=2, sticky="nw", padx=20, pady=(5, 10))

def add_status_row(row, label_text, variable):
    tk.Label(
        status_frame,
        text=label_text,
        bg=BG_COLOR,
        fg="#39FF0F"
    ).grid(row=row, column=0, sticky="w", padx=5, pady=3)

    tk.Label(
        status_frame,
        textvariable=variable,
        bg=BG_COLOR,
        fg="#39FF0F",
        anchor="w",
        justify="left",
        wraplength=700
    ).grid(row=row, column=1, sticky="w", padx=10, pady=3)

add_status_row(0, "Database:", database_path_var)
add_status_row(1, "Total Records:", total_records_var)
add_status_row(2, "Mail List Records:", mail_records_var)
add_status_row(3, "Email List Records:", email_records_var)
add_status_row(4, "Last Update:", last_update_var)


def refresh_database_status(update_timestamp=False):
    db_path = os.path.abspath(DB_FILE)

    with sqlite3.connect(DB_FILE) as conn:
        total_count = conn.execute(
            "SELECT COUNT(*) FROM recipients"
        ).fetchone()[0]

        mail_count = conn.execute(
            "SELECT COUNT(*) FROM recipients WHERE mail_list = 1"
        ).fetchone()[0]

        email_count = conn.execute(
            "SELECT COUNT(*) FROM recipients WHERE email_list = 1"
        ).fetchone()[0]

    database_path_var.set(db_path)
    total_records_var.set(str(total_count))
    mail_records_var.set(str(mail_count))
    email_records_var.set(str(email_count))

    if update_timestamp:
        last_update_var.set(save_last_update())
    else:
        last_update_var.set(get_saved_last_update())


# -------------------------
# DELETE BUTTON POSITION
# -------------------------

delete_frame = tk.Frame(root, bg=BG_COLOR)
delete_frame.grid(row=4, column=1, columnspan=2, sticky="n", padx=20, pady=(10, 20))

delete_button = ColorButton(
    delete_frame,
    text="Delete Record",
    state="disabled",
    bg=DELETE_RECORD_BG,
    fg="yellow",
    activebackground=DELETE_RECORD_ACTIVE_BG,
    activeforeground="DELETE_RECORD_ACTIVE_FG",
    highlightbackground=DELETE_RECORD_BG,
    command=delete_record
)
style_button(
    delete_button,
    bg=DELETE_RECORD_BG,
    fg="yellow",
    active_bg=DELETE_RECORD_ACTIVE_BG,
    active_fg=DELETE_RECORD_ACTIVE_FG,
    width=22
)
delete_button.grid(row=0, column=0, pady=5)


# -------------------------
# RIGHT-SIDE CONTROL FRAMES
# -------------------------

right_frame = tk.Frame(root, bg=BG_COLOR)
right_frame.grid(row=0, column=3, rowspan=4, sticky="ne", padx=20, pady=20)


edit_frame = tk.LabelFrame(
    right_frame,
    text="ID to Edit or Delete",
    bg=BG_COLOR,
    fg=FG_COLOR,
    padx=10,
    pady=10
)
edit_frame.grid(row=0, column=0, sticky="ew", pady=(0, 15))

edit_id_entry = tk.Entry(edit_frame, width=28)
edit_id_entry.grid(row=0, column=0, sticky="ew", padx=5, pady=5)

edit_id_entry.bind(
    "<Return>",
    lambda event: load_record_for_edit()
)

edit_id_entry.bind(
    "<KP_Enter>",
    lambda event: load_record_for_edit()
)

load_button = ColorButton(
    edit_frame,
    text="Load Record",
    command=load_record_for_edit
)
style_button(
    load_button,
    bg=SHOW_ALL_RECORDS_BG,
    fg=SHOW_ALL_RECORDS_FG,
    active_bg=SHOW_ALL_RECORDS_ACTIVE_BG,
    active_fg=SHOW_ALL_RECORDS_FG,
    width=28
)
load_button.grid(row=1, column=0, sticky="ew", padx=5, pady=5)

excel_frame = tk.LabelFrame(
    right_frame,
    text="Export to Excel",
    bg=BG_COLOR,
    fg=FG_COLOR,
    padx=10,
    pady=10
)
excel_frame.grid(row=1, column=0, sticky="ew", pady=(0, 15))

export_all_button = ColorButton(
    excel_frame,
    text="Export All (Excel)",
    command=lambda: export_xlsx(
        """
        SELECT
            id,
            honorific_title,
            first_name,
            last_name,
            area_code,
            phone,
            email,
            mailaddress_line1,
            mailaddress_line2,
            city,
            state,
            zip_code,
            email_list,
            mail_list,
            notes
        FROM recipients
        ORDER BY lower(last_name), lower(first_name)
        """,
        "all_recipients.xlsx"
    )
)
style_button(
    export_all_button,
    bg=EXPORT_BG,
    fg="white",
    active_bg=EXPORT_ACTIVE_BG,
    active_fg="white",
    width=28
)
export_all_button.grid(row=0, column=0, sticky="ew", padx=5, pady=5)

export_email_button = ColorButton(
    excel_frame,
    text="Export Email List (Excel)",
    command=lambda: export_xlsx(
        """
        SELECT honorific_title, first_name, last_name,
            email, area_code, phone, notes
        FROM recipients
        WHERE email_list = 1
        ORDER BY lower(last_name), lower(first_name)
        """,
        "email_list.xlsx"
    )
)
style_button(
    export_email_button,
    bg=EXPORT_BG,
    fg="white",
    active_bg=EXPORT_ACTIVE_BG,
    active_fg="white",
    width=28
)
export_email_button.grid(row=1, column=0, sticky="ew", padx=5, pady=5)

export_mail_button = ColorButton(
    excel_frame,
    text="Export Mail List (Excel)",
    command=lambda: export_xlsx(
        """
        SELECT
            id,
            honorific_title,
            first_name,
            last_name,
            area_code,
            phone,
            email,
            mailaddress_line1,
            mailaddress_line2,
            city,
            state,
            zip_code,
            email_list,
            mail_list,
            notes
        FROM recipients
        WHERE mail_list = 1
        ORDER BY lower(last_name), lower(first_name)
        """,
        "mail_list.xlsx"
    )
)
style_button(
    export_mail_button,
    bg=EXPORT_BG,
    fg="white",
    active_bg=EXPORT_ACTIVE_BG,
    active_fg="white",
    width=28
)
export_mail_button.grid(row=2, column=0, sticky="ew", padx=5, pady=5)


labels_frame = tk.LabelFrame(
    right_frame,
    text="Export Labels",
    bg=BG_COLOR,
    fg=FG_COLOR,
    padx=10,
    pady=10
)
labels_frame.grid(row=2, column=0, sticky="ew", pady=(0, 15))

export_labels_button = ColorButton(
    labels_frame,
    text='14 Labels 1.75" x 4"',
    command=export_mail_labels_5162
)
style_button(
    export_labels_button,
    bg=EXPORT_BG,
    active_bg=EXPORT_ACTIVE_BG,
    active_fg="white",
    fg="white",
    width=28
)
export_labels_button.grid(row=0, column=0, sticky="ew", padx=5, pady=5)

export_labels_5160_button = ColorButton(
    labels_frame,
    text='30 Labels 2.625" x 1"',
    command=export_mail_labels_5160
)
style_button(
    export_labels_5160_button,
    bg=EXPORT_BG,
    fg="white",
    active_bg=EXPORT_ACTIVE_BG,
    active_fg="white",

    width=28
)
export_labels_5160_button.grid(row=1, column=0, sticky="ew", padx=5, pady=5)


csv_frame = tk.LabelFrame(
    right_frame,
    text="Export / Import CSV",
    bg=BG_COLOR,
    fg=FG_COLOR,
    padx=10,
    pady=10
)
csv_frame.grid(row=3, column=0, sticky="ew", pady=(0, 15))

export_csv_button = ColorButton(
    csv_frame,
    text="Export All (CSV)",
    command=export_csv_all
)
style_button(
    export_csv_button,
    bg=EXPORT_BG,
    fg="white",
    active_bg=EXPORT_ACTIVE_BG,
    active_fg="white",
    width=28
)
export_csv_button.grid(row=0, column=0, sticky="ew", padx=5, pady=5)

import_csv_button = ColorButton(
    csv_frame,
    text="Import CSV",
    command=import_new_records_from_csv
)
style_button(
    import_csv_button,
    bg=EXPORT_BG,
    fg="white",
    active_bg=EXPORT_ACTIVE_BG,
    active_fg="white",
    width=28
)
import_csv_button.grid(row=1, column=0, sticky="ew", padx=5, pady=5)


refresh_database_status()

root.update_idletasks()

window_width = root.winfo_reqwidth()
screen_width = root.winfo_screenwidth()

x_position = int((screen_width - window_width) / 5)
y_position = 50

root.geometry(f"+{x_position}+{y_position}")


def handle_close_request():
    if not SESSION_HAS_UNBACKED_CHANGES:
        root.destroy()
        return

    dialog = tk.Toplevel(root)
    dialog.withdraw()
    dialog.title("CSV Backup Reminder")
    dialog.configure(bg=DUPLICATE_DIALOG_BG)
    dialog.transient(root)
    dialog.resizable(False, False)

    tk.Label(
        dialog,
        text=(
            "Changes have been made but not backed up.\n\n"
            "Would you like to create a CSV backup file now?"
        ),
        bg=DUPLICATE_DIALOG_BG,
        fg=FG_COLOR,
        justify="left",
        padx=20,
        pady=20
    ).grid(
        row=0,
        column=0,
        columnspan=2,
        sticky="w"
    )

    def export_before_close():
        dialog.destroy()

        # Closing is intentionally aborted.
        # After export, the user must close the app again.
        export_csv_all()

    def close_without_backup():
        dialog.destroy()
        root.destroy()

    ColorButton(
        dialog,
        text="Export CSV File",
        command=export_before_close,
        bg=EXPORT_BG,
        fg=EXPORT_FG,
        activebackground=EXPORT_ACTIVE_BG,
        activeforeground=EXPORT_FG,
        width=20
    ).grid(
        row=1,
        column=0,
        padx=(20, 10),
        pady=(0, 20)
    )

    ColorButton(
        dialog,
        text="Close Without Backup",
        command=close_without_backup,
        bg=DELETE_RECORD_BG,
        fg=DELETE_RECORD_FG,
        activebackground=DELETE_RECORD_ACTIVE_BG,
        activeforeground="black",
        width=20
    ).grid(
        row=1,
        column=1,
        padx=(10, 20),
        pady=(0, 20)
    )

    # Clicking the dialog's red close button cancels the close request
    # and returns the user to the application.
    dialog.protocol("WM_DELETE_WINDOW", dialog.destroy)
    center_toplevel_over_parent(dialog, root)
    dialog.grab_set()

# -------------------------
# LOGO IN BOTTOM-RIGHT CORNER
# -------------------------

logo_image = tk.PhotoImage(
    file=resource_path("Newsletter_logo.png")
)

logo_label = tk.Label(
    root,
    image=logo_image,
    bg=BG_COLOR,
    borderwidth=0,
    highlightthickness=0
)

logo_label.grid(
    row=4,
    column=3,
    sticky="se",
    padx=20,
    pady=(10, 20)
)


root.protocol("WM_DELETE_WINDOW", handle_close_request)

if root.tk.call("tk", "windowingsystem") == "aqua":
    root.createcommand("handle_close_request", handle_close_request)
    root.tk.eval("""
        proc ::tk::mac::Quit {} {
            handle_close_request
        }
    """)
    

root.mainloop()